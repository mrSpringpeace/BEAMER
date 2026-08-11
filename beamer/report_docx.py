# -*- coding: utf-8 -*-
"""Výpočtový protokol (DOCX) – tisknutelný dokument pro průkaz.

Zrcadlí textový protokol (`report.build_report`), ale strukturovaně: nadpisy,
tabulky a vložené grafy (schéma, VVÚ, posouzení). Volající (GUI) může předat
PNG obrázky svých pláten v `images` = {"schema": bytes, "vvu": bytes,
"margin": bytes}; když chybí, sekce s obrázkem se vynechá.
"""
from __future__ import annotations

import datetime
import io

from .i18n import tr
from .settings import fmt


def _tbl(doc, headers):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = str(h)
    return t


def _row(t, values):
    cells = t.add_row().cells
    for j, v in enumerate(values):
        cells[j].text = v if isinstance(v, str) else fmt(v)


def _img(doc, images, key, width_mm=160):
    if not images or key not in images or not images[key]:
        return
    try:
        from docx.shared import Mm
        doc.add_picture(io.BytesIO(images[key]), width=Mm(width_mm))
    except Exception:
        pass


def build_docx(state, result, margins, path, images=None, project_name=""):
    """Sestaví a uloží protokol DOCX. Vrací cestu."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from .sections_along import (normalized_segments, eff_defs,
                                 material_for_segment)
    from .section import build_section

    doc = Document()
    h = doc.add_heading(tr("BEAMER – Protokol statické analýzy nosníku"), 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(datetime.datetime.now().strftime("%Y-%m-%d %H:%M"))
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if project_name:
        pn = doc.add_paragraph(project_name)
        pn.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 1 Nosník a model ──
    doc.add_heading(tr("1  Nosník a model"), 1)
    doc.add_paragraph(f"{tr('Délka L')} = {fmt(state.length)} mm    "
                      f"{tr('Teorie ohybu')} = {state.theory}    "
                      f"{tr('Teorie torze')} = "
                      f"{getattr(state, 'torsion_theory', 'saint-venant')}    "
                      f"{tr('Dodatečný součinitel')} = {fmt(state.additional_factor)}")
    try:
        comb = state.active_combination()
        if comb:
            doc.add_paragraph(tr("Zobrazená kombinace: ") + comb.name)
    except Exception:
        pass

    doc.add_heading(tr("Podpory"), 2)
    support_headers = ["#", "x [mm]", tr("typ"), tr("úhel [°]"), tr("k / Δ")]
    vlasov = getattr(state, "torsion_theory", "saint-venant") == "vlasov"
    if vlasov:
        support_headers.append(tr("deplanace"))
    t = _tbl(doc, support_headers)
    for i, s in enumerate(state.supports):
        extra = ""
        if s.type == "spring":
            extra = f"k_z={fmt(getattr(s,'spring_z',0))} N/mm"
        elif abs(getattr(s, "settlement", 0.0)) > 1e-9:
            extra = f"Δ={fmt(s.settlement)} mm"
        row = [str(i + 1), fmt(s.x), s.type, fmt(s.angle), extra]
        if vlasov:
            hold = getattr(s, "restrain_warping", None)
            hold = (s.type == "fixed") if hold is None else bool(hold)
            row.append(tr("bráněná") if hold else tr("volná"))
        _row(t, row)

    doc.add_heading(tr("Zatížení"), 2)
    t = _tbl(doc, [tr("typ"), tr("popis")])
    from .report import _load_desc
    for ld in state.loads:
        _row(t, [ld.type, _load_desc(ld)])

    _img(doc, images, "schema")

    # ── 2 Úseky a průřezy ──
    doc.add_heading(tr("2  Úseky a průřezy"), 1)
    t = _tbl(doc, [tr("Úsek"), "x [mm]", tr("materiál"), tr("průřez"),
                   "A [mm²]", "Iy [mm⁴]", "Iz [mm⁴]", "IT [mm⁴]", "Iω [mm⁶]"])
    for i, seg in enumerate(normalized_segments(state)):
        sec1, _ = eff_defs(state, seg)
        mat = material_for_segment(state, seg)
        try:
            exact = (getattr(state, "tau_mode", "conservative") == "exact"
                     or getattr(state, "torsion_theory", "saint-venant") == "vlasov")
            sc = build_section(sec1, fem=exact, exact=exact)
            avals = [fmt(sc.A), fmt(sc.Iy), fmt(sc.Iz), fmt(sc.IT), fmt(sc.Iw)]
        except Exception:
            avals = ["—", "—", "—", "—", "—"]
        _row(t, [str(i + 1), f"{fmt(seg.x1)}–{fmt(seg.x2)}",
                 getattr(mat, "name", "?"), getattr(sec1, "type", "?")] + avals)

    # ── 3 Vnitřní účinky + reakce ──
    if result and result.is_stable and result.points:
        P = result.points
        doc.add_heading(tr("3  Vnitřní účinky (extrémy) a reakce"), 1)
        t = _tbl(doc, [tr("veličina"), tr("min"), tr("max"), tr("jednotka")])
        quantities = [("N", "N"), ("V", "N"), ("V_y", "N"),
                         ("M", "N·mm"), ("M_z", "N·mm"), ("Mk", "N·mm"),
                         ("w", "mm"), ("v", "mm"), ("phi_z", "rad")]
        if getattr(state, "torsion_theory", "saint-venant") == "vlasov":
            quantities += [("T_sv", "N·mm"), ("T_w", "N·mm"),
                           ("B", "N·mm²"), ("warping_rate", "1/mm")]
        for a, unit in quantities:
            vals = [getattr(p, a) for p in P]
            _row(t, [a, fmt(min(vals)), fmt(max(vals)), unit])
        _img(doc, images, "vvu")

        doc.add_heading(tr("Reakce"), 2)
        headers = ["x [mm]", "Rx [N]", "Ry [N]", "Rz [N]",
                   "My [N·mm]", "Mz [N·mm]", "Mx [N·mm]"]
        if getattr(state, "torsion_theory", "saint-venant") == "vlasov":
            headers.append("B [N·mm²]")
        t = _tbl(doc, headers)
        for rc in result.reactions:
            row = [fmt(rc.x), fmt(rc.Rx), fmt(rc.Ry_force), fmt(rc.Rz),
                   fmt(rc.Ry), fmt(rc.Rz_moment), fmt(rc.Rx_torsion)]
            if getattr(state, "torsion_theory", "saint-venant") == "vlasov":
                row.append(fmt(getattr(rc, "B_warping", 0.0)))
            _row(t, row)

    # ── 4 Posouzení (RF) ──
    if margins:
        crit = min(margins, key=lambda mm: mm.RF)
        doc.add_heading(tr("4  Posouzení (RF = reserve factor, ≥ 1 vyhovuje)"), 1)
        doc.add_paragraph(
            f"σ_red,max = {fmt(max(mm.mises_max for mm in margins))} MPa    "
            f"RF_min = {fmt(crit.RF)} ({crit.critical}) @ x = {fmt(crit.x)} mm")
        if crit.RF_local_buckling is not None:
            doc.add_paragraph(
                f"RF_local = {fmt(crit.RF_local_buckling)} "
                f"(stěna: {crit.critical_wall})")
        if crit.RF_crippling is not None:
            doc.add_paragraph(f"RF_crippling = {fmt(crit.RF_crippling)}")
        doc.add_paragraph(tr(
            "Lokální boulení: klasická desková teorie; crippling: empirická "
            "metoda Needham/Gerard, jen pro známou topologii."
        ))
        _img(doc, images, "margin")

    # ── Vzpěr, konzervativní kontrola (sdílená obálka přes kombinace) ──
    if result and result.is_stable and result.points:
        from .analysis import buckling_check, conservative_check, envelope_over_combinations
        env = None
        if getattr(state, "load_combinations", None):
            try:
                env = envelope_over_combinations(state)
            except Exception:
                env = None
        try:
            bc = buckling_check(state, result, env=env)
        except Exception:
            bc = None
        if bc is not None:
            scope = (tr("obálka přes všechny kombinace") if env is not None
                     else tr("zobrazená kombinace"))
            doc.add_heading(tr("Vzpěrná stabilita (Johnson-Euler, tlačené úseky)")
                            + f" – {scope}", 1)
            t = _tbl(doc, [tr("Úsek"), "N [N]", "λ", "σ_cr [MPa]", "P_cr [N]", "RF_vzpěr"])
            for r in bc.rows:
                _row(t, [r["label"], fmt(r["N"]), fmt(r["lam"]),
                         fmt(r["sigma_cr"]), fmt(r["P_cr"]), fmt(r["RF"])])
            doc.add_paragraph(f"RF_vzpěr,min = {fmt(bc.rf_min)} ({tr('řídí')} {bc.crit_label})")

        # fáze 2: bifurkace soustavy (vlastní čísla)
        from .analysis import buckling_eigen_check
        try:
            be = buckling_eigen_check(state, result)
        except Exception:
            be = None
        if be is not None:
            doc.add_heading(tr("Vzpěrná stabilita – fáze 2 (bifurkace soustavy, vlastní čísla)")
                            + f" – {tr('zobrazená kombinace')}", 1)
            doc.add_paragraph(f"λ_cr = {fmt(be.lam_cr)} = RF_vzpěr   "
                              f"P_cr = {fmt(be.P_cr)} N   N_ref = {fmt(be.N_ref)} N   "
                              f"μ_eff = {fmt(be.mu_eff)}")
            doc.add_paragraph(tr("Vzpěrná délka vyplývá z okrajových podmínek (bez ručního "
                                 "μ); slabá osa I_min; osové pole z rovnováhy jedné kombinace."))
            try:
                be_s = buckling_eigen_check(state, result, axis="max")
            except Exception:
                be_s = None
            if be_s is not None and abs(be_s.lam_cr - be.lam_cr) > 1e-3 * be.lam_cr:
                doc.add_paragraph(f"{tr('silná osa I_max')}: λ_cr = {fmt(be_s.lam_cr)}   "
                                  f"P_cr = {fmt(be_s.P_cr)} N   μ_eff = {fmt(be_s.mu_eff)}   "
                                  f"({tr('řídí slabá osa')})")
            _img(doc, images, "buckling")

        # beam-column: interakce tlak + ohyb (Bruhn)
        from .analysis import beam_column_check
        try:
            bcx = beam_column_check(state, result)
        except Exception:
            bcx = None
        if bcx is not None:
            doc.add_heading(tr("Interakce tlak + ohyb (beam-column, Bruhn)")
                            + f" – {tr('zobrazená kombinace')}", 1)
            t = _tbl(doc, [tr("Úsek"), "N [N]", "P_cr [N]", "R_c", "σ_ohyb [MPa]",
                           "R_b", "RF", "MS"])
            for r in bcx.rows:
                _row(t, [r["label"], fmt(r["N"]), fmt(r["P_cr"]), fmt(r["R_c"]),
                         fmt(r["sigma_b"]), fmt(r["R_b"]), fmt(r["RF"]), fmt(r["MS"])])
            doc.add_paragraph(f"RF_min = {fmt(bcx.rf_min)} ({tr('řídí')} {bcx.crit_label}) — "
                              + tr("R_c + R_b/(1−R_c) ≤ 1; RF = 1/R_int, MS = RF−1."))

        if env is not None:
            try:
                cc = conservative_check(state, env=env)
            except Exception:
                cc = None
            if cc is not None:
                doc.add_heading(tr("Konzervativní obálková kontrola (maxima naráz v řezu)"), 1)
                doc.add_paragraph(f"|N|={fmt(cc.N_max)} N   |V|={fmt(cc.V_max)} N   "
                                  f"|M|={fmt(cc.M_max)} N·mm   |Mk|={fmt(cc.Mk_max)} N·mm")
                t = _tbl(doc, [tr("Úsek"), "σ [MPa]", "τ [MPa]", "σ_red [MPa]", "RF"])
                for r in cc.rows:
                    _row(t, [r["label"], fmt(r["sigma"]), fmt(r["tau"]),
                             fmt(r["sred"]), fmt(r["RF"])])
                doc.add_paragraph(f"RF_konzervativní = {fmt(cc.rf_min)} "
                                  f"({tr('řídí')} {cc.crit_label})")

    doc.save(path)
    return path
