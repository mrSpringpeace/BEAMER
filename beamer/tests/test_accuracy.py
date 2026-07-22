"""Přesnostní (accuracy) testy – proti tabulkovým/analytickým hodnotám.

Doplněk ke konzistenčním testům v test_verification.py. Tyto testy by
zachytily chyby typu „IT obdélníku s prohozenými stranami" nebo „špatný
vzorec τ_t pro uzavřené průřezy" (audit v1.13 → opravy v1.14).
"""
from __future__ import annotations

import math
import os
import sys

import pytest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

from beamer.model import (
    Material, Support, Hinge, Load, LoadCase, LoadCombination,
    CrossSectionDef, SectionSegment, ProjectState, Body, Property,
)
from beamer.section import build_section
from beamer.solver import solve_beam
from beamer.analysis import forces_from_beam

MAT = Material("m_steel", "Steel", E=210000.0, G=81000.0, nu=0.3,
               rho=7.85, Re=235.0, Rm=360.0)


def make_state(length, supports, loads, theory="euler-bernoulli", hinges=None,
               sec=None):
    sec = sec or CrossSectionDef(type="rectangle", params={"b": 100.0, "h": 200.0})
    return ProjectState(
        length=float(length), supports=supports, hinges=hinges or [],
        load_cases=[LoadCase("lc", "LC", False)],
        load_combinations=[LoadCombination("c", "C", {"lc": 1.0})],
        loads=loads, materials=[Material(**vars(MAT))],
        selected_material_id=MAT.id, cross_section=sec,
        section_segments=[SectionSegment(0.0, float(length), sec, None,
                                         material_id=MAT.id)],
        additional_factor=1.0, theory=theory,
        selected_active_combination_id="c",
    )


def _rel(a, b):
    return abs(a - b) / (abs(b) if abs(b) > 1e-12 else 1.0)


def _tau_t_max(cs, Mk_nmm, n=120):
    f = forces_from_beam(N=0, V=0, M=0, Mk=Mk_nmm)
    prof = cs.profile(f, N=n)
    vals = [abs(p["tauT"]) for p in prof if not math.isnan(p["tauT"])]
    return max(vals) / 1e6 if vals else 0.0     # MPa


# ═══════════════════════════════════════════════════════════════
#  CHARAKTERISTIKY – přesnost proti analytice
# ═══════════════════════════════════════════════════════════════

def test_props_circle():
    D = 100.0
    cs = build_section(CrossSectionDef(type="circle", params={"D": D}))
    assert _rel(cs.A, math.pi * D**2 / 4) < 5e-3
    assert _rel(cs.Iy, math.pi * D**4 / 64) < 5e-3
    assert _rel(cs.IT, math.pi * D**4 / 32) < 5e-3


def test_props_tube():
    Do, t = 100.0, 5.0
    Di = Do - 2 * t
    cs = build_section(CrossSectionDef(type="tube", params={"Do": Do, "t": t}))
    assert _rel(cs.A, math.pi * (Do**2 - Di**2) / 4) < 5e-3
    assert _rel(cs.Iy, math.pi * (Do**4 - Di**4) / 64) < 5e-3
    assert _rel(cs.IT, math.pi * (Do**4 - Di**4) / 32) < 1e-6   # analyticky


def test_IT_rectangle_saint_venant():
    """C2 regrese: J = c1·a·t³ (kratší strana v kubíku), Roark/Timoshenko."""
    for b, h in [(100.0, 60.0), (100.0, 200.0), (80.0, 100.0)]:
        cs = build_section(CrossSectionDef(type="rectangle", params={"b": b, "h": h}))
        a_, t_ = max(b, h), min(b, h)
        c1 = 1/3 * (1 - 0.63*(t_/a_) + 0.052*(t_/a_)**5)
        J_ref = c1 * a_ * t_**3
        assert _rel(cs.IT, J_ref) < 0.02, f"IT {b}x{h}: {cs.IT} vs {J_ref}"


def test_IT_I_section_open_thin_walled():
    """I-profil: IT = Σ(1/3)·b·t³ (otevřený tenkostěnný)."""
    h, bf, tw, tf = 200.0, 100.0, 6.0, 10.0
    cs = build_section(CrossSectionDef(
        type="i_section",
        params={"h": h, "bf1": bf, "bf2": bf, "tw": tw, "tf1": tf, "tf2": tf}))
    hw = h - 2 * tf
    J_ref = (hw * tw**3 + 2 * bf * tf**3) / 3
    assert _rel(cs.IT, J_ref) < 1e-6


def test_composite_two_flanges():
    """Kompozit (2 pásnice bez stojiny): A, Iy přesně dle steineru."""
    bt = Body(points=[{"y": -50, "z": 80}, {"y": 50, "z": 80},
                      {"y": 50, "z": 100}, {"y": -50, "z": 100}])
    bb = Body(points=[{"y": -50, "z": -100}, {"y": 50, "z": -100},
                      {"y": 50, "z": -80}, {"y": -50, "z": -80}])
    cs = build_section(CrossSectionDef(type="polygon", bodies=[bt, bb]), fem=False)
    assert _rel(cs.A, 4000.0) < 1e-9
    Iy_ref = 2 * (100 * 20**3 / 12 + 100 * 20 * 90**2)
    assert _rel(cs.Iy, Iy_ref) < 1e-9
    assert abs(cs.cx) < 1e-9 and abs(cs.cz) < 1e-9


def test_square_with_hole():
    b = Body(points=[{"y": -50, "z": -50}, {"y": 50, "z": -50},
                     {"y": 50, "z": 50}, {"y": -50, "z": 50}],
             holes=[[{"y": -20, "z": -20}, {"y": 20, "z": -20},
                     {"y": 20, "z": 20}, {"y": -20, "z": 20}]])
    cs = build_section(CrossSectionDef(type="polygon", bodies=[b]), fem=False)
    assert _rel(cs.A, 8400.0) < 1e-9
    assert _rel(cs.Iy, (100**4 - 40**4) / 12) < 1e-9


def test_alpha_pl_rectangle():
    cs = build_section(CrossSectionDef(type="rectangle", params={"b": 100, "h": 200}))
    assert _rel(cs.alpha_pl, 1.5) < 0.01     # Wpl/Wel = 1.5 pro obdélník


# ═══════════════════════════════════════════════════════════════
#  TORZNÍ SMYKOVÉ NAPĚTÍ – per model (C1 regrese)
# ═══════════════════════════════════════════════════════════════

MK = 1.0e6      # N·mm


def test_tau_t_circle():
    D = 100.0
    cs = build_section(CrossSectionDef(type="circle", params={"D": D}))
    tau_ref = MK * (D/2) / cs.IT                       # MPa (mm jednotky)
    assert _rel(_tau_t_max(cs, MK), tau_ref) < 0.02


def test_tau_t_tube():
    Do, t = 100.0, 5.0
    cs = build_section(CrossSectionDef(type="tube", params={"Do": Do, "t": t}))
    tau_ref = MK * (Do/2) / cs.IT
    assert _rel(_tau_t_max(cs, MK), tau_ref) < 0.02


def test_tau_t_box_bredt():
    H, B, tw = 200.0, 100.0, 6.0
    cs = build_section(CrossSectionDef(type="box", params={"H": H, "B": B, "tw": tw}))
    Am = (H - tw) * (B - tw)
    tau_ref = MK / (2 * Am * tw)
    assert _rel(_tau_t_max(cs, MK), tau_ref) < 0.05


def test_tau_t_open_I():
    h, bf, tw, tf = 200.0, 100.0, 6.0, 10.0
    cs = build_section(CrossSectionDef(
        type="i_section",
        params={"h": h, "bf1": bf, "bf2": bf, "tw": tw, "tf1": tf, "tf2": tf}))
    tau_ref = MK * tf / cs.IT          # max na nejtlustší stěně (pásnice)
    assert _rel(_tau_t_max(cs, MK), tau_ref) < 0.10


# ═══════════════════════════════════════════════════════════════
#  NAPĚTÍ – velikost (ne jen znaménko)
# ═══════════════════════════════════════════════════════════════

def test_sigma_M_over_W():
    b, h, M = 100.0, 200.0, 5.0e6        # N·mm
    cs = build_section(CrossSectionDef(type="rectangle", params={"b": b, "h": h}))
    f = forces_from_beam(N=0, V=0, M=M, Mk=0)
    s = cs.stress(f, cs.z_top * 0.9999)
    W = b * h**2 / 6
    assert _rel(abs(s["sigma"]) / 1e6, M / W) < 5e-3


def test_tau_V_rectangle_parabola():
    """Žuravskij: τ_max = 1.5·V/A v neutrální ose obdélníku."""
    b, h, V = 100.0, 200.0, 1.0e5        # N
    cs = build_section(CrossSectionDef(type="rectangle", params={"b": b, "h": h}))
    f = forces_from_beam(N=0, V=V, M=0, Mk=0)
    s = cs.stress(f, 0.0)
    assert _rel(abs(s["tauVz"]) / 1e6, 1.5 * V / (b * h)) < 5e-3


# ═══════════════════════════════════════════════════════════════
#  SOLVER – další případy
# ═══════════════════════════════════════════════════════════════

def test_point_moment_reactions():
    L, M0 = 2000.0, 1.0e6
    st = make_state(L, [Support("a", 0, "pin", 0), Support("b", L, "roller", 0)],
                    [Load("m", "moment", "M", "lc", x=L/2, My=M0)])
    r = solve_beam(st)
    assert r.is_stable
    Rz = sorted(rc.Rz for rc in r.reactions)
    assert _rel(Rz[1], M0 / L) < 1e-6 and _rel(-Rz[0], M0 / L) < 1e-6
    assert abs(sum(rc.Rz for rc in r.reactions)) < 1e-6 * M0 / L


def test_hinge_zero_moment():
    L = 2000.0
    st = make_state(L, [Support("a", 0, "pin", 0), Support("b", L, "fixed", 0)],
                    [Load("f", "point_force", "F", "lc", x=500.0, Fz=-1000.0)],
                    hinges=[Hinge("h1", 1000.0)])
    r = solve_beam(st)
    assert r.is_stable
    p = min(r.points, key=lambda p: abs(p.x - 1000.0))
    Mmax = max(abs(q.M) for q in r.points)
    assert abs(p.M) < 1e-6 * max(Mmax, 1.0)
    assert _rel(sum(rc.Rz for rc in r.reactions), 1000.0) < 1e-9


def test_equilibrium_udl():
    L, q = 2000.0, 2.5
    st = make_state(L, [Support("a", 0, "pin", 0), Support("b", L, "roller", 0)],
                    [Load("q", "distributed", "q", "lc", x1=0, x2=L, q1=-q, q2=-q)])
    r = solve_beam(st)
    assert _rel(sum(rc.Rz for rc in r.reactions), q * L) < 1e-9


def test_skew_roller_45deg():
    """Šikmá rolna 45°: reakce musí ležet ve směru normály (Rx = Rz)
    a platí globální rovnováha. (M1 – dříve se úhel tiše ignoroval.)"""
    L, F = 2000.0, 1000.0
    st = make_state(L, [Support("a", 0, "pin", 0), Support("b", L, "roller", 45.0)],
                    [Load("f", "point_force", "F", "lc", x=L/2, Fz=-F)])
    r = solve_beam(st)
    assert r.is_stable
    rb = next(rc for rc in r.reactions if rc.support_type == "roller")
    ra = next(rc for rc in r.reactions if rc.support_type == "pin")
    assert _rel(rb.Rx, rb.Rz) < 1e-3                  # R ∥ n = (sin45, cos45)
    assert abs(ra.Rx + rb.Rx) < 1e-6 * F              # ΣFx = 0
    assert _rel(ra.Rz + rb.Rz, F) < 1e-6              # ΣFz = F


def test_timoshenko_udl_exact():
    """Timoshenko prostý nosník + UDL: δ_mid = 5qL⁴/384EI + qL²/(8·G·As)."""
    L, q = 2000.0, 1.0
    st = make_state(L, [Support("a", 0, "pin", 0), Support("b", L, "roller", 0)],
                    [Load("q", "distributed", "q", "lc", x1=0, x2=L, q1=-q, q2=-q)],
                    theory="timoshenko")
    r = solve_beam(st)
    cs = build_section(CrossSectionDef(type="rectangle", params={"b": 100, "h": 200}))
    As = cs.Asz if cs.Asz > 0 else cs.kappa * cs.A
    w_ref = 5*q*L**4/(384*MAT.E*cs.Iy) + q*L**2/(8*MAT.G*As)
    w_max = max(abs(p.w) for p in r.points)
    assert _rel(w_max, w_ref) < 5e-3


def test_unstable_beam_returns_message():
    """Nedostatečné podepření vrátí korektní chybovou hlášku (ne výjimku)."""
    L = 1000.0
    st = make_state(L, [], [Load("f", "point_force", "F", "lc", x=L/2, Fz=-1.0)])
    r = solve_beam(st)
    assert not r.is_stable
    assert r.error_message


# ═══════════════════════════════════════════════════════════════
#  BEAM – další analytické benchmarky (průhyb, reakce, i staticky neurčité)
# ═══════════════════════════════════════════════════════════════

def _EI():
    cs = build_section(CrossSectionDef(type="rectangle", params={"b": 100.0, "h": 200.0}))
    return MAT.E * cs.Iy


def test_cantilever_end_moment():
    """Konzola s koncovým momentem M: w_konec = M·L²/(2EI), φ_konec = M·L/EI,
    a ohybový moment po celé délce konstantní = M."""
    L, M = 1000.0, 1.0e6
    st = make_state(L, [Support("a", 0, "fixed", 0)],
                    [Load("m", "moment", "M", "lc", x=L, My=M)])
    r = solve_beam(st)
    assert r.is_stable
    EI = _EI()
    assert _rel(max(abs(p.w) for p in r.points), M*L*L/(2*EI)) < 1e-2
    assert _rel(max(abs(p.phi) for p in r.points), M*L/EI) < 1e-2
    assert _rel(max(abs(p.M) for p in r.points), M) < 1e-3


def test_simply_supported_central_point_load():
    """Prostý nosník, osamělá síla uprostřed P: w_mid = P·L³/(48EI), M_mid = P·L/4."""
    L, P = 2000.0, 1000.0
    st = make_state(L, [Support("a", 0, "pin", 0), Support("b", L, "roller", 0)],
                    [Load("p", "point_force", "P", "lc", x=L/2, Fz=-P)])
    r = solve_beam(st)
    assert r.is_stable
    EI = _EI()
    assert _rel(max(abs(p.w) for p in r.points), P*L**3/(48*EI)) < 1e-2
    assert _rel(max(abs(p.M) for p in r.points), P*L/4) < 1e-3


def test_propped_cantilever_udl():
    """Vetknutá–podepřená (staticky NEURČITÁ), spojité q: reakce v podpoře
    R = 3qL/8, moment u vetknutí = qL²/8 (a je to největší |M|)."""
    L, q = 2000.0, 2.0
    st = make_state(L, [Support("a", 0, "fixed", 0), Support("b", L, "roller", 0)],
                    [Load("q", "distributed", "q", "lc", x1=0, x2=L, q1=-q, q2=-q)])
    r = solve_beam(st)
    assert r.is_stable
    R_prop = next(rc for rc in r.reactions if rc.support_type == "roller")
    assert _rel(abs(R_prop.Rz), 3*q*L/8) < 5e-3
    assert _rel(max(abs(p.M) for p in r.points), q*L**2/8) < 5e-3


def test_fixed_fixed_udl():
    """Oboustranně vetknutá (staticky NEURČITÁ), spojité q: reakce qL/2 na
    každé straně, moment u konců qL²/12 (největší |M|)."""
    L, q = 2000.0, 2.0
    st = make_state(L, [Support("a", 0, "fixed", 0), Support("b", L, "fixed", 0)],
                    [Load("q", "distributed", "q", "lc", x1=0, x2=L, q1=-q, q2=-q)])
    r = solve_beam(st)
    assert r.is_stable
    for rc in r.reactions:
        assert _rel(abs(rc.Rz), q*L/2) < 5e-3
    assert _rel(max(abs(p.M) for p in r.points), q*L**2/12) < 5e-3


# ═══════════════════════════════════════════════════════════════
#  SLOŽENÝ PID (geometrické skládání z profilů knihovny)
# ═══════════════════════════════════════════════════════════════

def test_composite_pid_assembly():
    """Složený PID ze dvou knihovních obdélníků posunutých svisle → kompozit
    dvou těles; A = součet ploch, symetrie (Iyz≈0)."""
    from beamer.composite import composite_def
    from beamer.model import Property
    r1 = CrossSectionDef(type="rectangle", params={"b": 40, "h": 20}, id="s1", name="R1")
    r2 = CrossSectionDef(type="rectangle", params={"b": 40, "h": 20}, id="s2", name="R2")
    st = ProjectState(length=100.0, sections=[r1, r2],
                      materials=[Material(**vars(MAT))], selected_material_id=MAT.id)
    prop = Property(id="p1", pid=1, name="slozeny", composite_parts=[
        {"section_id": "s1", "dy": 0, "dz": 30, "angle": 0},
        {"section_id": "s2", "dy": 0, "dz": -30, "angle": 0},
    ])
    st.properties = [prop]
    cdef = composite_def(st, prop)
    assert cdef is not None and cdef.bodies is not None and len(cdef.bodies) == 2
    cs = build_section(cdef, fem=False)
    assert _rel(cs.A, 2 * 40 * 20) < 1e-6
    assert abs(cs.Iyz) < 1e-6 * cs.Iy            # svisle symetrické → Iyz ≈ 0
    # a přes eff_defs (úsek → PID) se vrátí týž složený průřez
    from beamer.sections_along import eff_defs
    seg = SectionSegment(0.0, 100.0, CrossSectionDef(), None, property_id="p1")
    st.section_segments = [seg]
    s1, s2 = eff_defs(st, seg)
    assert s2 is None and s1.bodies is not None and len(s1.bodies) == 2


def test_composite_tube_over_rod():
    """Trubka přes tyč (souosé): kompozit = mezikruží + plný kruh, A = obojí."""
    from beamer.composite import composite_def
    from beamer.model import Property
    rod = CrossSectionDef(type="circle", params={"D": 20}, id="rod", name="tyč")
    tube = CrossSectionDef(type="tube", params={"Do": 40, "t": 5}, id="tube", name="trubka")
    st = ProjectState(length=100.0, sections=[rod, tube],
                      materials=[Material(**vars(MAT))], selected_material_id=MAT.id)
    prop = Property(id="p", pid=1, name="tyč v trubce", composite_parts=[
        {"section_id": "rod", "dy": 0, "dz": 0, "angle": 0},
        {"section_id": "tube", "dy": 0, "dz": 0, "angle": 0},
    ])
    cs = build_section(composite_def(st, prop), fem=False)
    A_rod = math.pi * 10**2
    A_tube = math.pi * (20**2 - 15**2)
    assert _rel(cs.A, A_rod + A_tube) < 0.01     # kruhy polygonálně aprox.


def test_composite_weighted_reduces_to_geometric():
    """Jeden materiál: modulem vážené EIy = E · geometrické Iy (redukce), NA = těžiště."""
    from beamer.composite import composite_def, composite_weighted
    from beamer.model import Property
    r = CrossSectionDef(type="rectangle", params={"b": 40, "h": 20}, id="s", name="R")
    st = ProjectState(length=100.0, sections=[r], materials=[Material(**vars(MAT))],
                      selected_material_id=MAT.id)
    prop = Property(id="p", pid=1, name="c", composite_parts=[
        {"section_id": "s", "material_id": MAT.id, "dy": 0, "dz": 30, "angle": 0},
        {"section_id": "s", "material_id": MAT.id, "dy": 0, "dz": -30, "angle": 0},
    ])
    w = composite_weighted(st, prop)
    cs = build_section(composite_def(st, prop), fem=False)   # geometrický kompozit
    assert not w.multi_material
    assert abs(w.z_NA) < 1e-6 and abs(w.y_NA) < 1e-6         # symetrie
    assert _rel(w.EIy, MAT.E * cs.Iy) < 1e-6                 # EIy = E·Iy
    assert _rel(w.A_eq, cs.A) < 1e-6                         # A_eq = A (E_ref=E)


def test_composite_weighted_neutral_axis_shift():
    """Dva materiály (tužší nahoře): neutrální osa se posune k tuhému materiálu."""
    from beamer.composite import composite_weighted
    from beamer.model import Property
    steel = Material("st", "Ocel", E=210000.0, G=81000.0, nu=0.3, rho=7.85, Re=235, Rm=360)
    alu = Material("al", "Hliník", E=70000.0, G=27000.0, nu=0.33, rho=2.7, Re=200, Rm=300)
    r = CrossSectionDef(type="rectangle", params={"b": 40, "h": 20}, id="s", name="R")
    st = ProjectState(length=100.0, sections=[r], materials=[steel, alu],
                      selected_material_id="st")
    prop = Property(id="p", pid=1, name="c", composite_parts=[
        {"section_id": "s", "material_id": "st", "dy": 0, "dz": 30, "angle": 0},   # tuhá nahoře
        {"section_id": "s", "material_id": "al", "dy": 0, "dz": -30, "angle": 0},
    ])
    w = composite_weighted(st, prop)
    assert w.multi_material
    A = 40 * 20
    # z_NA = ΣEᵢAᵢzᵢ / ΣEᵢAᵢ = A·30·(Est−Eal)/(A·(Est+Eal))
    z_expect = 30 * (210000 - 70000) / (210000 + 70000)      # = 15.0 mm nahoru
    assert _rel(w.z_NA, z_expect) < 1e-6
    assert w.z_NA > 0                                        # posun k tužšímu (nahoru)
    assert _rel(w.EA, (210000 + 70000) * A) < 1e-6


def test_composite_beam_deflection_uses_weighted_EI():
    """Prostý nosník s osamělou silou uprostřed, průřez = složený PID ze DVOU
    materiálů (ocel nahoře, hliník dole). Průhyb musí odpovídat modulem vážené
    tuhosti: w_mid = P·L³/(48·EIy). Ověřuje napojení B1 do solveru."""
    from beamer.model import Property
    from beamer.composite import composite_weighted, composite_def
    steel = Material("st", "Ocel", E=210000, G=81000, nu=0.3, rho=7.85, Re=235, Rm=360)
    alu = Material("al", "Hliník", E=70000, G=27000, nu=0.33, rho=2.7, Re=200, Rm=300)
    r = CrossSectionDef(type="rectangle", params={"b": 40, "h": 20}, id="s", name="R")
    prop = Property(id="p", pid=1, name="bimetal", composite_parts=[
        {"section_id": "s", "material_id": "st", "dy": 0, "dz": 10, "angle": 0},
        {"section_id": "s", "material_id": "al", "dy": 0, "dz": -10, "angle": 0},
    ])
    L, P = 2000.0, 1000.0
    seg = SectionSegment(0.0, L, CrossSectionDef(), None, property_id="p")
    st = ProjectState(
        length=L, supports=[Support("a", 0, "pin", 0), Support("b", L, "roller", 0)],
        load_cases=[LoadCase("lc", "LC", False)],
        load_combinations=[LoadCombination("c", "C", {"f": 1.0})],
        loads=[Load("f", "point_force", "P", "lc", x=L/2, Fz=-P)],
        materials=[steel, alu], selected_material_id="st", sections=[r],
        properties=[prop], cross_section=CrossSectionDef(type="rectangle", params={"b": 40, "h": 20}),
        section_segments=[seg], additional_factor=1.0, selected_active_combination_id="c")
    w = composite_weighted(st, prop)
    res = solve_beam(st)
    assert res.is_stable
    w_max = max(abs(p.w) for p in res.points)
    assert _rel(w_max, P*L**3/(48*w.EIy)) < 1e-2       # průhyb dle vážené EIy
    assert _rel(max(abs(p.M) for p in res.points), P*L/4) < 1e-3   # M nezávisí na materiálu
    # směs ocel+hliník je poddajnější než celá ocel (kontrola, že vážení působí)
    EIy_all_steel = 210000 * build_section(composite_def(st, prop), fem=False).Iy
    assert w.EIy < EIy_all_steel


def test_composite_assessment_wired():
    """Posouzení složeného nosníku vrací per-materiálové RF (values_at_x i
    reserves_along_beam), řídicí RF = min přes materiály."""
    from beamer.model import Property
    from beamer.analysis import values_at_x_multi, reserves_along_beam
    steel = Material("st", "Ocel", E=210000, G=81000, nu=0.3, rho=7.85, Re=235, Rm=360)
    alu = Material("al", "Hliník", E=70000, G=27000, nu=0.33, rho=2.7, Re=200, Rm=300)
    r = CrossSectionDef(type="rectangle", params={"b": 40, "h": 20}, id="s", name="R")
    seg = SectionSegment(0.0, 2000.0, CrossSectionDef(), None, property_id="p")
    st = ProjectState(
        length=2000.0, supports=[Support("a", 0, "pin", 0), Support("b", 2000, "roller", 0)],
        load_cases=[LoadCase("lc", "LC", False)],
        load_combinations=[LoadCombination("c", "C", {"f": 1.0})],
        loads=[Load("f", "point_force", "P", "lc", x=1000, Fz=-5000.0)],
        materials=[steel, alu], selected_material_id="st", sections=[r],
        properties=[Property(id="p", pid=1, name="bimetal", composite_parts=[
            {"section_id": "s", "material_id": "st", "dy": 0, "dz": 10, "angle": 0},
            {"section_id": "s", "material_id": "al", "dy": 0, "dz": -10, "angle": 0}])],
        cross_section=CrossSectionDef(type="rectangle", params={"b": 40, "h": 20}),
        section_segments=[seg], additional_factor=1.0, selected_active_combination_id="c")
    res = solve_beam(st)
    d = values_at_x_multi(res, st, 1000.0)[0]
    assert d.get("composite") and len(d.get("materials", [])) == 2
    assert d["RF"] == min(m["RF_yield"] for m in d["materials"]) or \
           d["RF"] == min(min(m["RF_yield"], m["RF_ultimate"]) for m in d["materials"])
    marg = reserves_along_beam(res, st)
    assert marg and all(m.RF > 0 for m in marg)


def test_bimetal_thermal_self_stress():
    """Bimetal (ocel+hliník, různá α) při rovnoměrném ΔT: i bez vnějšího zatížení
    vzniká samovyrovnané vnitřní pnutí (N=0, M=0), protože vrstvy se chtějí
    roztáhnout různě. Ověřuje EAα/ESα i vzorec σᵢ=Eᵢ(ε₀+κ·dz−αᵢΔT) proti ruční
    referenci a self-equilibraci (ΣσᵢAᵢ=0)."""
    from beamer.model import Property
    from beamer.composite import composite_weighted, composite_stress
    steel = Material("st", "Ocel", E=210000, G=81000, nu=0.3, rho=7.85,
                     Re=235, Rm=360, alpha=12e-6)
    alu = Material("al", "Hliník", E=70000, G=27000, nu=0.33, rho=2.7,
                   Re=200, Rm=300, alpha=23e-6)
    r = CrossSectionDef(type="rectangle", params={"b": 40, "h": 20}, id="s", name="R")
    prop = Property(id="p", pid=1, name="bimetal", composite_parts=[
        {"section_id": "s", "material_id": "st", "dy": 0, "dz": 10, "angle": 0},
        {"section_id": "s", "material_id": "al", "dy": 0, "dz": -10, "angle": 0}])
    st = ProjectState(length=100.0, sections=[r], materials=[steel, alu],
                      selected_material_id="st", properties=[prop])
    w = composite_weighted(st, prop)
    # sekční charakteristiky proti ruční referenci
    assert _rel(w.z_NA, 5.0) < 1e-9
    assert _rel(w.EAalpha, 3304.0) < 1e-9          # ΣEᵢαᵢAᵢ
    assert _rel(w.ESalpha, -9240.0) < 1e-9         # ΣEᵢαᵢAᵢ(zᵢ−z_NA)

    dT = 50.0
    eps0 = dT * w.EAalpha / w.EA
    kth = dT * w.ESalpha / w.EIy
    ref = {"Ocel": 210000*abs(eps0 + kth*(0-w.z_NA) - 12e-6*dT),
           "Hliník": 70000*abs(eps0 + kth*(0-w.z_NA) - 23e-6*dT)}
    rows = composite_stress(st, prop, N=0.0, M=0.0, dT=dT)
    by = {x["material"]: x for x in rows}
    assert _rel(by["Ocel"]["sigma_max"], ref["Ocel"]) < 1e-3
    assert _rel(by["Hliník"]["sigma_max"], ref["Hliník"]) < 1e-3
    assert by["Ocel"]["sigma_max"] > 40.0          # nezanedbatelné pnutí
    # bez ΔT žádné samovyrovnané pnutí
    rows0 = composite_stress(st, prop, N=0.0, M=0.0, dT=0.0)
    assert max(x["sigma_max"] for x in rows0) < 1e-6


def test_bimetal_thermal_free_bar_end_to_end():
    """Volně uložený (staticky určitý) bimetalový nosník při ΔT: reakce/N/M≈0,
    ale posouzení podél nosníku nese nenulové samovyrovnané pnutí v obou
    materiálech (protažení celým řetězcem solver→analýza→kompozit)."""
    from beamer.model import Property
    from beamer.analysis import values_at_x_multi
    steel = Material("st", "Ocel", E=210000, G=81000, nu=0.3, rho=7.85,
                     Re=235, Rm=360, alpha=12e-6)
    alu = Material("al", "Hliník", E=70000, G=27000, nu=0.33, rho=2.7,
                   Re=200, Rm=300, alpha=23e-6)
    r = CrossSectionDef(type="rectangle", params={"b": 40, "h": 20}, id="s", name="R")
    L = 2000.0
    seg = SectionSegment(0.0, L, CrossSectionDef(), None, property_id="p")
    th = Load("t", "thermal", "dT", "lc"); th.x1 = 0; th.x2 = L; th.dT = 50.0
    st = ProjectState(
        length=L, supports=[Support("a", 0, "pin", 0), Support("b", L, "roller", 0)],
        load_cases=[LoadCase("lc", "LC", False)],
        load_combinations=[LoadCombination("c", "C", {"t": 1.0})],
        loads=[th],
        materials=[steel, alu], selected_material_id="st", sections=[r],
        properties=[Property(id="p", pid=1, name="bimetal", composite_parts=[
            {"section_id": "s", "material_id": "st", "dy": 0, "dz": 10, "angle": 0},
            {"section_id": "s", "material_id": "al", "dy": 0, "dz": -10, "angle": 0}])],
        cross_section=CrossSectionDef(type="rectangle", params={"b": 40, "h": 20}),
        section_segments=[seg], additional_factor=1.0, selected_active_combination_id="c")
    res = solve_beam(st)
    assert res.is_stable
    # staticky určitý → vnitřní síly ≈ 0 (volná dilatace)
    assert max(abs(p.N) for p in res.points) < 1e-3
    assert max(abs(p.M) for p in res.points) < 1e-3
    d = values_at_x_multi(res, st, L/2)[0]
    assert d.get("composite")
    smax = {m["material"]: m["sigma_max"] for m in d["materials"]}
    assert smax["Ocel"] > 40.0 and smax["Hliník"] > 15.0   # self-stress se objeví


def test_composite_torsion_end_to_end():
    """Vetknutý složený nosník + kroutící moment: per-materiálové posouzení
    projde celým řetězcem (solver GJ_eff → analýza) a nese torzní τ > 0 →
    σ_red > σ v obou materiálech."""
    from beamer.model import Property
    from beamer.analysis import values_at_x_multi
    steel = Material("st", "Ocel", E=210000, G=81000, nu=0.3, rho=7.85, Re=235, Rm=360)
    alu = Material("al", "Hliník", E=70000, G=27000, nu=0.33, rho=2.7, Re=200, Rm=300)
    r = CrossSectionDef(type="rectangle", params={"b": 40, "h": 20}, id="s", name="R")
    seg = SectionSegment(0.0, 1000.0, CrossSectionDef(), None, property_id="p")
    st = ProjectState(
        length=1000.0, supports=[Support("a", 0, "fixed", 0)],
        load_cases=[LoadCase("lc", "LC", False)],
        load_combinations=[LoadCombination("c", "C", {"t": 1.0})],
        loads=[Load("t", "torsion", "T", "lc", x=1000, Mx=5.0e5)],
        materials=[steel, alu], selected_material_id="st", sections=[r],
        properties=[Property(id="p", pid=1, name="bimetal", composite_parts=[
            {"section_id": "s", "material_id": "st", "dy": 0, "dz": 10, "angle": 0},
            {"section_id": "s", "material_id": "al", "dy": 0, "dz": -10, "angle": 0}])],
        cross_section=CrossSectionDef(type="rectangle", params={"b": 40, "h": 20}),
        section_segments=[seg], additional_factor=1.0, selected_active_combination_id="c")
    res = solve_beam(st)
    d = values_at_x_multi(res, st, 500.0)[0]
    assert d.get("composite") and abs(d["Mk"]) > 1.0
    assert d["tau_max"] > 0.0
    for mm in d["materials"]:
        assert mm["tau_max"] > 0.0
        assert mm["mises_max"] >= mm["sigma_max"]


def test_composite_stress_per_material():
    """Napětí per materiál (bimetal pod momentem): tužší materiál (ocel) nese
    vyšší napětí; RF se počítá z jeho vlastní pevnosti."""
    from beamer.model import Property
    from beamer.composite import composite_stress
    steel = Material("st", "Ocel", E=210000, G=81000, nu=0.3, rho=7.85, Re=235, Rm=360)
    alu = Material("al", "Hliník", E=70000, G=27000, nu=0.33, rho=2.7, Re=200, Rm=300)
    r = CrossSectionDef(type="rectangle", params={"b": 40, "h": 20}, id="s", name="R")
    st = ProjectState(length=100.0, sections=[r], materials=[steel, alu],
                      selected_material_id="st")
    prop = Property(id="p", pid=1, name="bimetal", composite_parts=[
        {"section_id": "s", "material_id": "st", "dy": 0, "dz": 10, "angle": 0},
        {"section_id": "s", "material_id": "al", "dy": 0, "dz": -10, "angle": 0},
    ])
    rows = composite_stress(st, prop, N=0.0, M=1.0e6)
    assert rows is not None and len(rows) == 2
    by = {r_["material"]: r_ for r_ in rows}
    assert by["Ocel"]["sigma_max"] > by["Hliník"]["sigma_max"]     # tužší nese víc
    # RF z vlastní pevnosti materiálu
    assert by["Ocel"]["RF_yield"] == steel.Re / by["Ocel"]["sigma_max"]


def test_section_rotation_90_swaps_Iy_Iz():
    """Natočení průřezu o 90° prohodí Iy a Iz (plocha beze změny)."""
    r0 = build_section(CrossSectionDef(type="rectangle", params={"b": 40, "h": 20}), fem=False)
    r90 = build_section(CrossSectionDef(type="rectangle", params={"b": 40, "h": 20}, rotation=90), fem=False)
    assert _rel(r90.A, r0.A) < 1e-9
    assert _rel(r90.Iy, r0.Iz) < 1e-3
    assert _rel(r90.Iz, r0.Iy) < 1e-3


def test_section_rotation_45_transformed_inertia():
    """Natočení o 45°: Iy=Iz=(Iy0+Iz0)/2 a |Iyz|=|Iy0−Iz0|/2 (Mohrova transformace)."""
    r0 = build_section(CrossSectionDef(type="rectangle", params={"b": 40, "h": 20}), fem=False)
    r45 = build_section(CrossSectionDef(type="rectangle", params={"b": 40, "h": 20}, rotation=45), fem=False)
    assert _rel(r45.Iy, (r0.Iy + r0.Iz)/2) < 1e-3
    assert _rel(r45.Iz, (r0.Iy + r0.Iz)/2) < 1e-3
    assert _rel(abs(r45.Iyz), abs(r0.Iy - r0.Iz)/2) < 1e-3


def test_pid_rotation_applied_and_beam():
    """Natočení na PID (90°) se projeví v průřezu i v průhybu nosníku
    (loading se přepočítá na natočený profil = transformované Iy)."""
    from beamer.model import Property
    r = CrossSectionDef(type="rectangle", params={"b": 40, "h": 20}, id="s", name="R")
    seg = SectionSegment(0.0, 2000.0, CrossSectionDef(), None, property_id="p")
    st = ProjectState(
        length=2000.0, supports=[Support("a", 0, "pin", 0), Support("b", 2000, "roller", 0)],
        load_cases=[LoadCase("lc", "LC", False)],
        load_combinations=[LoadCombination("c", "C", {"f": 1.0})],
        loads=[Load("f", "point_force", "P", "lc", x=1000, Fz=-1000.0)],
        materials=[Material(**vars(MAT))], selected_material_id=MAT.id, sections=[r],
        properties=[Property(id="p", pid=1, name="X", sec1_id="s", rotation=90)],
        cross_section=CrossSectionDef(type="rectangle", params={"b": 40, "h": 20}),
        section_segments=[seg], additional_factor=1.0, selected_active_combination_id="c")
    from beamer.sections_along import eff_defs
    s1, _ = eff_defs(st, seg)
    cs = build_section(s1, fem=False)
    r0 = build_section(r, fem=False)
    assert _rel(cs.Iy, r0.Iz) < 1e-3          # PID 90° → Iy = Iz0
    res = solve_beam(st)
    assert res.is_stable
    w_max = max(abs(p.w) for p in res.points)
    assert _rel(w_max, 1000.0*2000.0**3/(48*MAT.E*cs.Iy)) < 1e-2   # průhyb dle natočené Iy


# ═══════════════════════════════════════════════════════════════
#  KŘÍŽOVÁ VALIDACE proti sectionproperties (nezávislá knihovna, dev-only)
# ═══════════════════════════════════════════════════════════════

def test_xval_rotation_vs_sectionproperties():
    """Natočený obdélník: Iy/Iz/Iyz musí sedět s sectionproperties."""
    import pytest
    pytest.importorskip("sectionproperties")
    from sectionproperties.pre.library import rectangular_section
    from sectionproperties.analysis import Section
    for ang in (0, 30, 45, 90):
        g = rectangular_section(d=20, b=40).rotate_section(angle=ang)
        g.create_mesh(mesh_sizes=[0.5])
        s = Section(g); s.calculate_geometric_properties()
        ixx, iyy, ixy = s.get_ic()
        cs = build_section(CrossSectionDef(type="rectangle", params={"b": 40, "h": 20},
                                           rotation=ang), fem=False)
        assert _rel(cs.Iy, ixx) < 1e-6
        assert _rel(cs.Iz, iyy) < 1e-6
        assert _rel(abs(cs.Iyz), abs(ixy)) < 1e-6 or abs(cs.Iyz) < 1e-3


def test_xval_composite_vs_sectionproperties():
    """Složený dvoumateriál: EA, EIy k neutrální ose a poloha NA vs
    sectionproperties (modulem vážené / transformované charakteristiky)."""
    import pytest
    pytest.importorskip("sectionproperties")
    from sectionproperties.pre.library import rectangular_section
    from sectionproperties.pre import Material as SPMat
    from sectionproperties.analysis import Section
    from beamer.model import Property
    from beamer.composite import composite_weighted
    st_sp = SPMat("st", 210000, 0.3, 235, 7.85e-6, "grey")
    al_sp = SPMat("al", 70000, 0.33, 200, 2.7e-6, "blue")
    geom = (rectangular_section(d=20, b=40, material=st_sp).shift_section(-20, 0)
            + rectangular_section(d=20, b=40, material=al_sp).shift_section(-20, -20))
    geom.create_mesh(mesh_sizes=[0.5])
    s = Section(geom); s.calculate_geometric_properties()
    EA_sp, EIy_sp, cz_sp = s.get_ea(), s.get_eic()[0], s.get_c()[1]
    steel = Material("st", "Ocel", E=210000, G=81000, nu=0.3, rho=7.85, Re=235, Rm=360)
    alu = Material("al", "Hliník", E=70000, G=27000, nu=0.33, rho=2.7, Re=200, Rm=300)
    r = CrossSectionDef(type="rectangle", params={"b": 40, "h": 20}, id="s", name="R")
    prop = Property(id="p", pid=1, name="c", composite_parts=[
        {"section_id": "s", "material_id": "st", "dy": 0, "dz": 10, "angle": 0},
        {"section_id": "s", "material_id": "al", "dy": 0, "dz": -10, "angle": 0}])
    stt = ProjectState(length=100, sections=[r], materials=[steel, alu], selected_material_id="st")
    w = composite_weighted(stt, prop)
    assert _rel(w.EA, EA_sp) < 1e-6
    assert _rel(w.EIy, EIy_sp) < 1e-6
    assert _rel(w.z_NA, cz_sp) < 1e-6


def test_envelope_over_combinations():
    """Obálka přes kombinace: dvě proporcionální kombinace (1× a 2×). Řídicí je
    2× (nižší RF); crit RF obálky = min RF přes kombinace; VVÚ obálka lineárně
    škáluje (M_max při 2× = 2·M_max při 1×), min RF obálky = polovina RF při 1×."""
    from beamer.analysis import envelope_over_combinations, reserves_along_beam
    seg = SectionSegment(0.0, 2000.0, CrossSectionDef(type="rectangle",
                         params={"b": 40, "h": 80}), None)
    st = ProjectState(
        length=2000.0, materials=[MAT], selected_material_id="m_steel",
        cross_section=CrossSectionDef(type="rectangle", params={"b": 40, "h": 80}),
        section_segments=[seg],
        supports=[Support("a", 0, "pin", 0), Support("b", 2000, "roller", 0)],
        load_cases=[LoadCase("lc", "LC", False)],
        loads=[Load("f", "point_force", "P", "lc", x=1000, Fz=-5000.0)],
        load_combinations=[
            LoadCombination(id="c1", name="1x", factors={"f": 1.0}),
            LoadCombination(id="c2", name="2x", factors={"f": 2.0})],
        additional_factor=1.0)
    env = envelope_over_combinations(st)
    assert env is not None and env.n_combos == 2
    assert env.crit_combo == "2x"                       # 2× řídí (nižší RF)
    # M obálka: max při 2× = 2× max při 1×
    from beamer.solver import solve_beam
    r1 = solve_beam(st, factors={"f": 1.0})
    m1 = max(abs(p.M) for p in r1.points)
    assert _rel(max(abs(min(env.M_min)), abs(max(env.M_max))), 2 * m1) < 1e-6
    # crit RF obálky = min RF přes kombinace
    rf1 = min(r.RF for r in reserves_along_beam(r1, st))
    assert _rel(env.crit_rf, rf1 / 2.0) < 5e-3          # RF ~ 1/zatížení


def test_load_curve_parse_and_generate():
    """Spojité z křivky: parser (komentáře, oddělovače, desetinná čárka) +
    po částech lineární segmenty; výslednice = lichoběžníková integrace."""
    from beamer.loadgen import parse_xq_curve, loads_from_curve
    txt = ("# hlavicka\n0 -2.0\n500  -3,5\n1000 ; -1.0\n2000\t0\n"
           "; komentar\n1000 -1.0\n")   # duplicitní x=1000 → sloučí
    pts = parse_xq_curve(txt)
    assert pts == [(0.0, -2.0), (500.0, -3.5), (1000.0, -1.0), (2000.0, 0.0)]
    loads = loads_from_curve(pts, "lc", "q")
    assert len(loads) == 3 and all(l.type == "distributed" for l in loads)
    assert loads[0].x1 == 0 and loads[0].q2 == -3.5 and loads[1].q1 == -3.5
    R = sum((l.q1 + l.q2) / 2 * (l.x2 - l.x1) for l in loads)
    assert _rel(R, -3000.0) < 1e-9      # ∫q dx přes lichoběžníky


def test_direct_section_independent_chars():
    """Direct profil: nezávislé A, Iy, Iz, IT se přenesou přesně; legacy (jen Iy)
    odvodí A ze čtverce (zpětná kompatibilita); solver bere EI (prohyb konzoly)."""
    d = CrossSectionDef(type="direct",
                        params={"A": 2500.0, "Iy": 5.2e6, "Iz": 1.3e6, "IT": 3.1e6})
    cs = build_section(d, fem=False)
    assert _rel(cs.A, 2500.0) < 1e-9 and _rel(cs.Iy, 5.2e6) < 1e-9
    assert _rel(cs.Iz, 1.3e6) < 1e-9 and _rel(cs.IT, 3.1e6) < 1e-9
    assert _rel(cs.iy, (5.2e6 / 2500.0) ** 0.5) < 1e-9
    # legacy: jen Iy → A ze čtverce h⁴/12=Iy
    cl = build_section(CrossSectionDef(type="direct", params={"Iy": 5.2e6}), fem=False)
    h = (12 * 5.2e6) ** 0.25
    assert _rel(cl.A, h * h) < 1e-6 and _rel(cl.Iy, 5.2e6) < 1e-6
    # solver: konzola P na konci → w = P·L³/(3·E·Iy)
    seg = SectionSegment(0.0, 1000.0, d, None)
    st = ProjectState(
        length=1000, materials=[MAT], selected_material_id="m_steel",
        cross_section=d, section_segments=[seg],
        supports=[Support("a", 0, "fixed", 0)],
        load_cases=[LoadCase("lc", "LC", False)],
        loads=[Load("f", "point_force", "P", "lc", x=1000, Fz=-1000.0)],
        load_combinations=[LoadCombination(id="c", name="c", factors={"f": 1.0})],
        additional_factor=1.0)
    res = solve_beam(st)
    w_tip = min(p.w for p in res.points)
    assert _rel(w_tip, -1000 * 1000**3 / (3 * 210000 * 5.2e6)) < 1e-6


def test_conservative_check_is_upper_bound():
    """Konzervativní obálková kontrola (G): σ_red z maxim ≥ přesné σ_red, tedy
    RF_konz ≤ RF_přesný. Maxima obálky = max |N|,|V|,|M|,|Mk| přes kombinace;
    dosazena naráz do kritického řezu (√(σ_max²+3τ_max²))."""
    from beamer.analysis import (conservative_check, envelope_over_combinations)
    seg = SectionSegment(0.0, 2000.0, CrossSectionDef(type="rectangle",
                         params={"b": 40, "h": 80}), None)
    st = ProjectState(
        length=2000.0, materials=[MAT], selected_material_id="m_steel",
        cross_section=CrossSectionDef(type="rectangle", params={"b": 40, "h": 80}),
        section_segments=[seg],
        supports=[Support("a", 0, "pin", 0), Support("b", 2000, "roller", 0)],
        load_cases=[LoadCase("lc", "LC", False)],
        loads=[Load("f", "point_force", "P", "lc", x=1000, Fz=-5000.0)],
        load_combinations=[
            LoadCombination(id="c1", name="1x", factors={"f": 1.0}),
            LoadCombination(id="c2", name="1.6x", factors={"f": 1.6})],
        additional_factor=1.0, rf_basis="yield")
    cc = conservative_check(st, basis="yield")
    env = envelope_over_combinations(st)
    assert cc is not None and cc.n_combos == 2
    # maxima = 1.6× kombinace
    assert _rel(cc.M_max, 1.6 * 5000.0 * 2000.0 / 4) < 1e-6      # prostý nosník: M=PL/4
    # ruční σ = M/W + N/A, W = b·h²/6; τ = 1.5 V/A; σ_red = √(σ²+3τ²)
    A = 40 * 80; W = 40 * 80**2 / 6
    sig = cc.M_max / W
    tau = 1.5 * cc.V_max / A
    import math
    assert _rel(cc.rows[0]["sred"], math.sqrt(sig**2 + 3 * tau**2)) < 2e-2
    # konzervativní RF ≤ přesný RF (horní odhad napětí)
    assert cc.rf_min <= env.crit_rf + 1e-6


def _ss_state(supports, loads=None):
    rect = CrossSectionDef(type="rectangle", params={"b": 40, "h": 80})
    return ProjectState(
        length=2000.0, materials=[MAT], selected_material_id="m_steel",
        cross_section=rect, section_segments=[SectionSegment(0.0, 2000.0, rect, None)],
        supports=supports, load_cases=[LoadCase("lc", "LC", False)],
        loads=loads or [], additional_factor=1.0)


def test_spring_support_stiff_limit_and_reaction():
    """Pružná podpora (8): velmi tuhá pružina = tuhá podpora; měkká pružina má
    reakci R = k·w (síla pružiny)."""
    P = Load("f", "point_force", "P", "lc", x=1000, Fz=-5000.0)
    r0 = solve_beam(_ss_state([Support("a", 0, "pin", 0), Support("b", 2000, "pin", 0)], [P]))
    w0 = min(p.w for p in r0.points)
    stiff = Support("b", 2000, "spring", 0); stiff.spring_z = 1e9
    r1 = solve_beam(_ss_state([Support("a", 0, "pin", 0), stiff], [P]))
    assert _rel(min(p.w for p in r1.points), w0) < 1e-4        # tuhá pružina = pin
    soft = Support("b", 2000, "spring", 0); soft.spring_z = 500.0
    r2 = solve_beam(_ss_state([Support("a", 0, "pin", 0), soft], [P]))
    w_supp = [p.w for p in r2.points if abs(p.x - 2000) < 1][0]
    Rz = [rc.Rz for rc in r2.reactions][-1]
    assert _rel(abs(Rz), abs(500.0 * w_supp)) < 1e-3           # reakce = k·w


def test_biaxial_iyz_coupling_rotation_equivalence():
    """Biaxiál fáze B2: Iyz tuhostní spřažení (šikmý ohyb). Natočit SEKCI o α a
    zatížit svisle = nechat sekci a natočit ZATÍŽENÍ o −α (stejná fyzika). Celková
    deformace |d|=√(w²+v²) i max napětí musí být invariantní → ověří správné
    znaménko/velikost křížové tuhosti EIyz i vyhodnocení napětí v pravých vláknech."""
    from beamer.analysis import reserves_along_beam
    E, b, h, L, P, a = 210000.0, 40.0, 80.0, 1000.0, 1000.0, 45.0
    matB = Material("mb", "St", E=E, G=81000.0, nu=0.3, rho=7.85, Re=235.0, Rm=360.0)

    def build(rot, Fy, Fz):
        rect = CrossSectionDef(type="rectangle", params={"b": b, "h": h}, rotation=rot)
        ld = Load("f", "point_force", "F", "lc", x=L/2, Fy=Fy, Fz=Fz)
        return ProjectState(
            length=L, materials=[matB], selected_material_id="mb",
            cross_section=rect, section_segments=[SectionSegment(0.0, L, rect, None)],
            supports=[Support("a", 0, "pin", 0), Support("b", L, "roller", 0)],
            load_cases=[LoadCase("lc", "LC", False)], loads=[ld],
            load_combinations=[LoadCombination(id="c", name="c", factors={"f": 1.0})],
            additional_factor=1.0)

    def solve(s):
        r = solve_beam(s)
        p = [q for q in r.points if abs(q.x - L/2) < 6][0]
        return math.hypot(p.w, p.v), max(x.sigma_max for x in reserves_along_beam(r, s))

    c = math.cos(math.radians(a))
    dA, sA = solve(build(a, 0.0, -P))          # sekce +45°, svislá síla
    dB, sB = solve(build(0.0, -P*c, -P*c))     # sekce 0°, síla natočená o −45°
    assert _rel(dA, dB) < 1e-3                  # celková deformace invariantní
    assert _rel(sA, sB) < 5e-3                  # max napětí invariantní
    # bez spřažení (Iyz=0) je v-rovina od svislé síly nulová (kontrola regrese)
    r0 = solve_beam(build(0.0, 0.0, -P))
    assert max(abs(p.v) for p in r0.points) < 1e-9


def test_biaxial_stress_into_rf():
    """Biaxiál fáze B1: současný ohyb v obou rovinách (Fz+Fy) → normálové napětí
    se sčítá v rohu průřezu: σ_max = |My|·z_ext/Iy + |Mz|·y_ext/Iz. Ověření přes
    celý pipeline (reserves_along_beam) i regrese Mz=0 = uniaxiál."""
    from beamer.analysis import reserves_along_beam
    E, b, h, L, P = 210000.0, 40.0, 80.0, 1000.0, 1000.0
    matB = Material("mb", "St", E=E, G=81000.0, nu=0.3, rho=7.85, Re=235.0, Rm=360.0)
    rect = CrossSectionDef(type="rectangle", params={"b": b, "h": h})
    Iy, Iz = b*h**3/12, h*b**3/12
    Mmid = P*L/4
    s_uni = Mmid*(h/2)/Iy                       # jen svislý ohyb
    s_bx = Mmid*(h/2)/Iy + Mmid*(b/2)/Iz        # oba ohyby v rohu

    def mk(*comps):
        loads = []
        for i, c in enumerate(comps):
            ld = Load(f"f{i}", "point_force", "F", "lc", x=L/2)
            setattr(ld, c, -P); loads.append(ld)
        return ProjectState(
            length=L, materials=[matB], selected_material_id="mb",
            cross_section=rect, section_segments=[SectionSegment(0.0, L, rect, None)],
            supports=[Support("a", 0, "pin", 0), Support("b", L, "roller", 0)],
            load_cases=[LoadCase("lc", "LC", False)], loads=loads,
            load_combinations=[LoadCombination(id="c", name="c",
                                               factors={l.id: 1.0 for l in loads})],
            additional_factor=1.0)

    # uniaxiální (jen Fz) a biaxiální (Fz+Fy) – poměr eliminuje vzorkování stanic
    r1 = reserves_along_beam(solve_beam(mk("Fz")), mk("Fz"))
    st2 = mk("Fz", "Fy")
    r2 = reserves_along_beam(solve_beam(st2), st2)
    sig1 = max(x.sigma_max for x in r1)
    sig2 = max(x.sigma_max for x in r2)
    assert _rel(sig2 / sig1, s_bx / s_uni) < 5e-3     # poměr napětí = 3.0 (roh)
    # RF kleslo ve stejném poměru
    rf1 = min(x.RF for x in r1); rf2 = min(x.RF for x in r2)
    assert _rel(rf1 / rf2, s_bx / s_uni) < 1e-2
    # uniaxiální absolutně (blízko midspanu) – biaxiální větev se nezapnula
    assert 0.98 * s_uni < sig1 < 1.001 * s_uni


def test_biaxial_horizontal_bending():
    """Biaxiál fáze A (6 DOF, dvě dekuplované roviny): vodorovná síla Fy dá ohyb
    v rovině x-y se slabou tuhostí EIz – zrcadlo svislého případu s Iz místo Iy.
    Prostě podepřený, síla uprostřed: v_mid = P·L³/(48·E·Iz), M_z,mid = P·L/4.
    Svislá rovina zůstane netknutá (w≈0, N≈0, Mk≈0)."""
    E, b, h, L, P = 210000.0, 40.0, 80.0, 1000.0, 1000.0
    matB = Material("mb", "St", E=E, G=81000.0, nu=0.3, rho=7.85, Re=235.0, Rm=360.0)
    rect = CrossSectionDef(type="rectangle", params={"b": b, "h": h})
    Iz = h * b**3 / 12.0                       # slabá osa (vodorovný ohyb)
    v_exp = P * L**3 / (48.0 * E * Iz)
    Mz_exp = P * L / 4.0

    def mk(comp):
        ld = Load("f", "point_force", "F", "lc", x=L/2)
        setattr(ld, comp, -P)
        return ProjectState(
            length=L, materials=[matB], selected_material_id="mb",
            cross_section=rect, section_segments=[SectionSegment(0.0, L, rect, None)],
            supports=[Support("a", 0, "pin", 0), Support("b", L, "roller", 0)],
            load_cases=[LoadCase("lc", "LC", False)], loads=[ld],
            load_combinations=[LoadCombination(id="c", name="c", factors={"f": 1.0})],
            additional_factor=1.0)

    r = solve_beam(mk("Fy"))
    v_mid = [p.v for p in r.points if abs(p.x - L/2) < 6][0]
    Mz_mid = max(abs(p.M_z) for p in r.points)
    assert _rel(abs(v_mid), v_exp) < 5e-3            # průhyb slabou osou
    assert _rel(Mz_mid, Mz_exp) < 1e-6               # M_z = P·L/4
    # svislá rovina netknutá
    assert max(abs(p.w) for p in r.points) < 1e-9
    assert max(abs(p.M) for p in r.points) < 1e-6
    assert max(abs(p.N) for p in r.points) < 1e-6

    # symetrie: Fy s Iz dá stejný |průhyb| jako Fz s Iy (zrcadlo)
    rz = solve_beam(mk("Fz"))
    Iy = b * h**3 / 12.0
    w_mid = [p.w for p in rz.points if abs(p.x - L/2) < 6][0]
    assert _rel(abs(v_mid) / abs(w_mid), Iy / Iz) < 5e-3   # poměr průhybů = Iy/Iz


def test_rotated_parametric_shear_center_warping():
    """Natočení parametrického profilu: Iω a Wk jsou rotačně invariantní, střed
    smyku je bod, který se otáčí s profilem. Dřív scanline aproximace běžela na
    natočené geometrii → špatný SC (y_SC natvrdo 0) a Iω. Oprava počítá z
    nenatočeného dvojníka a vektor SC otočí. Iy/Iz musí zůstat tenzorově otočené."""
    def mk(rot):
        return build_section(CrossSectionDef(type="l_section",
                             params={"h": 100, "b": 80, "t": 8}, rotation=rot))
    s0 = mk(0.0)
    assert abs(s0.z_SC) > 1.0            # L má SC mimo těžiště (jinak test nic netestuje)
    for rot in (30.0, 90.0, 180.0):
        s = mk(rot)
        th = math.radians(rot); ca, sa = math.cos(th), math.sin(th)
        y_exp = s0.y_SC * ca - s0.z_SC * sa
        z_exp = s0.y_SC * sa + s0.z_SC * ca
        assert math.hypot(s.y_SC - y_exp, s.z_SC - z_exp) < 1e-6   # SC = rotace SC(0)
        assert _rel(s.Iw, s0.Iw) < 1e-9                            # Iω invariantní
        assert _rel(s.Wk, s0.Wk) < 1e-9                            # Wk invariantní
    # Iy/Iz se natočením 90° prohodí (nedotčeno opravou SC)
    s90 = mk(90.0)
    assert _rel(s90.Iy, s0.Iz) < 1e-6 and _rel(s90.Iz, s0.Iy) < 1e-6


def test_schema_load_filter_by_combination():
    """Schéma kreslí zatížení v kontextu ZOBRAZENÉ kombinace: jen zatížení s
    nenulovým faktorem, velikost × faktor × dodatečný součinitel. Bez kombinace
    fallback = zadané hodnoty."""
    # kreslení potřebuje matplotlib + Qt binding (plots.py = QtAgg) → v CI bez GUI
    # stacku se přeskočí místo selhání
    pytest.importorskip("matplotlib")
    pytest.importorskip("PySide6")
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from beamer.gui.plots import _draw_schema

    rect = CrossSectionDef(type="rectangle", params={"b": 40, "h": 80})
    st = ProjectState(
        length=1000, materials=[MAT], selected_material_id="m_steel",
        cross_section=rect,
        supports=[Support("a", 0, "pin", 0), Support("b", 1000, "roller", 0)],
        load_cases=[LoadCase("lc1", "LC1"), LoadCase("lc2", "LC2")],
        loads=[Load("A", "point_force", "A", "lc1", x=300, Fz=-1000),
               Load("B", "point_force", "B", "lc2", x=700, Fz=-2000)],
        load_combinations=[LoadCombination("C1", "C1", {"A": 1.0}),
                           LoadCombination("C2", "C2", {"A": 1.0, "B": 0.5})],
        additional_factor=1.0)

    def force_labels(active, add=1.0):
        st.selected_active_combination_id = active
        st.additional_factor = add
        fig = plt.figure(); ax = fig.add_subplot(111)
        _draw_schema(ax, st, None)
        vals = [t.get_text() for t in ax.texts
                if " N" in t.get_text() and "N/mm" not in t.get_text()]
        plt.close(fig)
        return vals

    assert force_labels("C1") == ["F1\n-1000 N"]                    # jen A
    assert force_labels("C2") == ["F1\n-1000 N", "F2\n-1000 N"]     # A + B×0.5
    assert force_labels("C1", add=1.5) == ["F1\n-1500 N"]           # × dodatečný souč.

    # přepínač „napříč kombinacemi" (filter_by_combination=False) → vše raw i při
    # aktivní kombinaci C1
    st.selected_active_combination_id = "C1"; st.additional_factor = 1.0
    fig = plt.figure(); ax = fig.add_subplot(111)
    _draw_schema(ax, st, None, filter_by_combination=False)
    allv = [t.get_text() for t in ax.texts
            if " N" in t.get_text() and "N/mm" not in t.get_text()]
    plt.close(fig)
    assert allv == ["F1\n-1000 N", "F2\n-2000 N"]
    st.load_combinations = []; st.selected_active_combination_id = ""
    st.additional_factor = 1.0
    fig = plt.figure(); ax = fig.add_subplot(111)
    _draw_schema(ax, st, None)                                       # bez kombinace = raw
    raw = [t.get_text() for t in ax.texts
           if " N" in t.get_text() and "N/mm" not in t.get_text()]
    plt.close(fig)
    assert raw == ["F1\n-1000 N", "F2\n-2000 N"]


def test_support_gap_contact():
    """Vůle podpory (nelineární kontakt, aktivní množina): střední podpora s
    vůlí g nechá uzel volný v ±g. Malé zatížení (průhyb < g) → nedosedne
    (reakce 0, chová se jako 2-podporový nosník); velké (průhyb > g) → dosedne
    (w = −g, reakce > 0). g=0 = tuhá podpora (regrese)."""
    rect = CrossSectionDef(type="rectangle", params={"b": 40, "h": 80})

    def mk(P, gap):
        smid = Support("mid", 1000, "roller", 0); smid.gap = gap
        return ProjectState(
            length=2000, materials=[MAT], selected_material_id="m_steel",
            cross_section=rect, section_segments=[SectionSegment(0.0, 2000.0, rect, None)],
            supports=[Support("a", 0, "pin", 0), smid, Support("b", 2000, "roller", 0)],
            load_cases=[LoadCase("lc", "LC", False)],
            loads=[Load("f", "point_force", "P", "lc", x=1000, Fz=-P)],
            load_combinations=[LoadCombination(id="c", name="c", factors={"f": 1.0})],
            additional_factor=1.0)

    def wmid(res):
        return [p.w for p in res.points if abs(p.x - 1000) < 1][0]

    def rmid(res):
        return [rc.Rz for rc in res.reactions if abs(rc.x - 1000) < 1][0]

    # 2-podporový průhyb (bez střední): w = P·L³/(48EI); P=1000 dá |w|<1, P=5000 >1
    EI = 210000 * (40 * 80**3 / 12)
    w_2supp = 5000 * 2000**3 / (48 * EI)
    assert w_2supp > 1.0                       # velké zatížení překročí vůli 1 mm
    # malé zatížení: nedosedne
    rA = solve_beam(mk(1000, 1.0))
    assert abs(wmid(rA)) < 1.0 and abs(rmid(rA)) < 1.0
    # velké: dosedne přesně na vůli, reakce tlačí
    rB = solve_beam(mk(5000, 1.0))
    assert _rel(wmid(rB), -1.0) < 1e-6 and rmid(rB) > 100.0
    # regrese: g=0 == tuhá střední podpora (w_mid ≈ 0)
    assert abs(wmid(solve_beam(mk(5000, 0.0)))) < 1e-6


def test_support_settlement():
    """Předepsaný posun podpory (8): pravá podpora prostého nosníku klesne o Δ,
    bez zatížení → nosník se lineárně nakloní, w(0)=0, w(L)=Δ."""
    sb = Support("b", 2000, "pin", 0); sb.settlement = -3.0
    r = solve_beam(_ss_state([Support("a", 0, "pin", 0), sb], []))
    wa = [p.w for p in r.points if abs(p.x - 0) < 1][0]
    wm = [p.w for p in r.points if abs(p.x - 1000) < 1][0]
    wb = [p.w for p in r.points if abs(p.x - 2000) < 1][0]
    assert abs(wa) < 1e-6 and _rel(wm, -1.5) < 1e-3 and _rel(wb, -3.0) < 1e-3


def test_buckling_euler_and_johnson():
    """Vzpěr fáze 1 (A): štíhlý tlačený sloup (μ=1) → Euler P_cr=π²·E·I_min/L²,
    RF_vzpěr=P_cr/|N|; krátký sloup → Johnson (σ_cr<Euler, σ_cr≤Re)."""
    import math
    from beamer.analysis import buckling_check, _johnson_euler_sigma_cr
    rect = CrossSectionDef(type="rectangle", params={"b": 40, "h": 80})
    seg = SectionSegment(0.0, 2000.0, rect, None); seg.buckling_mu = 1.0
    st = ProjectState(
        length=2000, materials=[MAT], selected_material_id="m_steel",
        cross_section=rect, section_segments=[seg],
        supports=[Support("a", 0, "pin", 0), Support("b", 2000, "roller", 0)],
        load_cases=[LoadCase("lc", "LC", False)],
        loads=[Load("f", "point_force", "P", "lc", x=2000, Fx=-50000.0)],
        load_combinations=[LoadCombination(id="c", name="c", factors={"f": 1.0})],
        additional_factor=1.0)
    res = solve_beam(st)
    bc = buckling_check(st, res)
    assert bc is not None
    Iz = 80 * 40**3 / 12
    P_euler = math.pi**2 * 210000 * Iz / 2000**2
    assert _rel(bc.rows[0]["P_cr"], P_euler) < 1e-6
    assert _rel(bc.rf_min, P_euler / 50000.0) < 1e-6
    # Johnson: krátká štíhlost < λ_cr → σ_cr menší než Euler i než Re
    lam_short = 50.0
    lam_cr = math.pi * math.sqrt(2 * 210000 / 235)
    assert lam_short < lam_cr
    sig_j = _johnson_euler_sigma_cr(210000, 235, lam_short)
    assert sig_j < 235 and sig_j < math.pi**2 * 210000 / lam_short**2
    # A5: vzpěr přes OBÁLKU kombinací – přidaná 2× kombinace řídí (RF/2),
    # i když je zobrazená 1×
    from beamer.analysis import envelope_over_combinations
    st.load_combinations.append(LoadCombination(id="c2", name="2x", factors={"f": 2.0}))
    env = envelope_over_combinations(st)
    bc_env = buckling_check(st, res, env=env)
    assert _rel(bc_env.rf_min, bc.rf_min / 2.0) < 1e-6


def test_docx_report_builds():
    """Protokol DOCX (C): sestaví se, má nadpisy a tabulky (vstupy/úseky/reakce)."""
    import pytest, os, tempfile
    pytest.importorskip("docx")
    from beamer.report_docx import build_docx
    from beamer.analysis import reserves_along_beam
    seg = SectionSegment(0.0, 2000.0, CrossSectionDef(type="rectangle",
                         params={"b": 40, "h": 80}), None)
    st = ProjectState(
        length=2000.0, materials=[MAT], selected_material_id="m_steel",
        cross_section=CrossSectionDef(type="rectangle", params={"b": 40, "h": 80}),
        section_segments=[seg],
        supports=[Support("a", 0, "pin", 0), Support("b", 2000, "roller", 0)],
        load_cases=[LoadCase("lc", "LC", False)],
        loads=[Load("f", "point_force", "P", "lc", x=1000, Fz=-5000.0)],
        load_combinations=[LoadCombination(id="c", name="1x", factors={"f": 1.0})],
        selected_active_combination_id="c", additional_factor=1.0)
    res = solve_beam(st)
    rsv = reserves_along_beam(res, st)
    path = os.path.join(tempfile.gettempdir(), "beamer_test_report.docx")
    build_docx(st, res, rsv, path)
    from docx import Document
    d = Document(path)
    heads = [p.text for p in d.paragraphs if p.style.name.startswith("Heading")]
    assert any("Nosník" in h for h in heads) and any("Posouzení" in h for h in heads)
    assert len(d.tables) >= 3
    os.remove(path)


def test_resolver_rotated_pid_multi_segment():
    """Regrese (kritická): u nosníku s VÍCE úseky a NATOČENÝMI PID (rotation≠0)
    resolver dřív cachoval podle id() přechodné kopie definice → po recyklaci id
    (GC) dostal úsek průřez SOUSEDNÍHO úseku (špatné Iy → špatné VVÚ/RF,
    nedeterministicky). Test: každý úsek musí dát SVŮJ Iy."""
    from beamer.model import Property
    from beamer.sections_along import SectionResolver
    import gc
    secA = CrossSectionDef(type="rectangle", params={"b": 40, "h": 80}, id="sa", name="A")
    secB = CrossSectionDef(type="rectangle", params={"b": 60, "h": 120}, id="sb", name="B")
    pA = Property(id="pa", pid=1, name="A", sec1_id="sa", material_id="m_steel", rotation=90)
    pB = Property(id="pb", pid=2, name="B", sec1_id="sb", material_id="m_steel", rotation=90)
    seg1 = SectionSegment(0.0, 100.0, CrossSectionDef(), None, property_id="pa")
    seg2 = SectionSegment(100.0, 160.0, CrossSectionDef(), None, property_id="pb")
    st = ProjectState(length=160.0, materials=[MAT], selected_material_id="m_steel",
                      sections=[secA, secB], properties=[pA, pB],
                      section_segments=[seg1, seg2])
    # očekávané Iy po natočení 90° (prohodí b,h): rect(b,h) rot90 → Iy = h·b³/12
    IyA = 80 * 40**3 / 12
    IyB = 120 * 60**3 / 12
    r = SectionResolver(st)
    for x in range(2, 160, 3):
        gc.collect()                      # vynuť recyklaci id (spouštěč staré chyby)
        iy = r.at(float(x)).Iy
        expect = IyA if x < 100 else IyB
        assert _rel(iy, expect) < 1e-9, f"x={x}: Iy={iy} != {expect}"


def test_thermal_axial_load():
    """Teplotní zatížení (D): volná dilatace → N=0 (bez pnutí); plně osově vázaný
    nosník (fixed-fixed) při ohřevu ΔT → tlak N=−E·A·α·ΔT."""
    matT = Material("mt", "St", E=210000.0, G=81000.0, nu=0.3, rho=7.85,
                    Re=235.0, Rm=360.0, alpha=12e-6)
    rect = CrossSectionDef(type="rectangle", params={"b": 40, "h": 80})

    def mk(supports):
        th = Load("t", "thermal", "dT", "lc"); th.x1 = 0; th.x2 = 2000; th.dT = 50.0
        return ProjectState(
            length=2000, materials=[matT], selected_material_id="mt",
            cross_section=rect, section_segments=[SectionSegment(0.0, 2000.0, rect, None)],
            supports=supports, load_cases=[LoadCase("lc", "LC", False)], loads=[th],
            load_combinations=[LoadCombination(id="c", name="c", factors={"t": 1.0})],
            additional_factor=1.0)

    r_free = solve_beam(mk([Support("a", 0, "fixed", 0)]))
    assert abs(r_free.points[len(r_free.points) // 2].N) < 1e-3     # volný → bez pnutí
    r_fix = solve_beam(mk([Support("a", 0, "fixed", 0), Support("b", 2000, "fixed", 0)]))
    N = r_fix.points[len(r_fix.points) // 2].N
    assert _rel(N, -210000 * 3200 * 12e-6 * 50.0) < 1e-6           # tlak = −EA·α·ΔT
    assert N < 0                                                    # ohřev vázaný → tlak
    # A1: hranice teploty MIMO pravidelnou mřížku (x2=1025) musí být uzel →
    # ohřátá část fixed-fixed: N = −EA·α·ΔT·(L_teplá/L) přesně
    def mk2(x2):
        th = Load("t", "thermal", "dT", "lc"); th.x1 = 0; th.x2 = x2; th.dT = 50.0
        return ProjectState(
            length=2000, materials=[matT], selected_material_id="mt",
            cross_section=rect, section_segments=[SectionSegment(0.0, 2000.0, rect, None)],
            supports=[Support("a", 0, "fixed", 0), Support("b", 2000, "fixed", 0)],
            load_cases=[LoadCase("lc", "LC", False)], loads=[th],
            load_combinations=[LoadCombination(id="c", name="c", factors={"t": 1.0})],
            additional_factor=1.0)
    r2 = solve_beam(mk2(1025.0))
    N2 = r2.points[len(r2.points) // 2].N
    assert _rel(N2, -210000 * 3200 * 12e-6 * 50.0 * (1025 / 2000)) < 1e-6


def test_composite_biaxial_stress():
    """Kompozit biaxiálně (fáze D): svislé + vodorovné zatížení symetrického
    složeného průřezu (ocel. trubka + dural. jádro). Pro symetrický řez a stejné
    Fz=Fy je výsledný ohyb √2× → napětí v rohu = √2 × jednoosé. Ověří biaxiální
    člen v composite_stress_field. Uniaxiál (Fy=0) zůstává beze změny."""
    from beamer.analysis import reserves_along_beam
    stl = Material("st", "Ocel", E=210000, G=81000, nu=0.3, rho=7.85, Re=235, Rm=360)
    alu = Material("al", "Dural", E=72000, G=27000, nu=0.33, rho=2.78, Re=280, Rm=440)
    tube = CrossSectionDef(id="t", name="t", type="tube", params={"Do": 60, "t": 6})
    core = CrossSectionDef(id="c", name="c", type="circle", params={"D": 48})
    prop = Property(id="p", pid=1, name="k", material_id="st",
                   composite_parts=[{"section_id": "t", "material_id": "st", "dy": 0, "dz": 0, "angle": 0},
                                    {"section_id": "c", "material_id": "al", "dy": 0, "dz": 0, "angle": 0}])

    def mk(fy):
        seg = SectionSegment(0.0, 1000.0, tube, None); seg.property_id = "p"
        loads = [Load("fz", "point_force", "sv", "lc", x=500, Fz=-4000)]
        if fy:
            loads.append(Load("fy", "point_force", "vo", "lc", x=500, Fy=-4000))
        return ProjectState(
            length=1000.0, materials=[stl, alu], selected_material_id="st",
            sections=[tube, core], properties=[prop], section_segments=[seg],
            supports=[Support("a", 0, "pin"), Support("b", 1000, "roller")],
            load_cases=[LoadCase("lc", "LC")], loads=loads,
            load_combinations=[LoadCombination(id="cc", name="cc",
                               factors={l.id: 1.0 for l in loads})], additional_factor=1.0)

    st1 = mk(False); s1 = max(x.sigma_max for x in reserves_along_beam(solve_beam(st1), st1))
    st2 = mk(True); s2 = max(x.sigma_max for x in reserves_along_beam(solve_beam(st2), st2))
    assert _rel(s2 / s1, math.sqrt(2.0)) < 1e-2      # symetrický řez, Fz=Fy → √2×


def test_beam_column_bruhn_interaction():
    """Interakce tlak+ohyb (beam-column) leteckou interakční rovnicí Bruhna:
    R_c + R_b/(1−R_c) ≤ 1, RF = 1/R_int. Ověření proti ručnímu výpočtu na
    tlačené vzpěře s boční silou (sprajc)."""
    from beamer.analysis import beam_column_check
    from beamer.section import build_section
    matB = Material("mb", "Dural", E=72000.0, G=27000.0, nu=0.33, rho=2.78,
                    Re=280.0, Rm=440.0)
    tube = CrossSectionDef(type="tube", params={"Do": 40, "t": 3})
    P, Plat, L = 5000.0, 200.0, 1000.0
    st = ProjectState(
        length=L, materials=[matB], selected_material_id="mb", cross_section=tube,
        section_segments=[SectionSegment(0.0, L, tube, None)],
        supports=[Support("a", 0, "pin"), Support("b", L, "roller")],
        load_cases=[LoadCase("lc", "LC", False)],
        loads=[Load("fx", "point_force", "tlak", "lc", x=L, Fx=-P),
               Load("fz", "point_force", "bocni", "lc", x=L/2, Fz=-Plat)],
        load_combinations=[LoadCombination(id="c", name="c",
                           factors={"fx": 1.0, "fz": 1.0})], additional_factor=1.0)
    bc = beam_column_check(st, solve_beam(st))
    assert bc is not None and len(bc.rows) == 1
    row = bc.rows[0]

    # ruční Bruhn
    sec = build_section(tube, fem=False)
    P_cr = math.pi**2 * 72000.0 * sec.Iy / L**2          # Euler, μ=1
    R_c = P / P_cr
    sig_b = (Plat*L/4) * 20.0 / sec.Iy                   # M·c/I, c=Do/2
    R_b = sig_b / min(280.0, 440.0)
    R_int = R_c + R_b / (1.0 - R_c)
    assert _rel(row["R_c"], R_c) < 5e-3
    assert _rel(row["R_b"], R_b) < 5e-3
    assert _rel(row["RF"], 1.0/R_int) < 5e-3
    assert _rel(row["MS"], 1.0/R_int - 1.0) < 5e-3
    # čistý tah → žádný beam-column (None)
    st.loads[0].Fx = +P
    assert beam_column_check(st, solve_beam(st)) is None


def test_buckling_eigen_euler_cases():
    """Vzpěr fáze 2 (vlastní čísla): kritický násobitel λ_cr celé soustavy pro
    tři klasické Eulerovy případy → odvozený součinitel vzpěrné délky μ musí
    vyjít 1.0 (kloub-kloub), 0.5 (vetknutý-vetknutý), 2.0 (vetknutý-volný).
    Tím se ověří geometrická matice tuhosti i okrajové podmínky (w=0 vs w=0&φ=0)."""
    from beamer.analysis import buckling_eigen_check
    E, b, h, L, P = 210000.0, 40.0, 80.0, 1000.0, 1000.0
    matB = Material("mb", "St", E=E, G=81000.0, nu=0.3, rho=7.85,
                    Re=235.0, Rm=360.0, alpha=12e-6)
    rect = CrossSectionDef(type="rectangle", params={"b": b, "h": h})
    Imin = min(b*h**3/12, h*b**3/12)
    P_euler = math.pi**2 * E * Imin / L**2

    def run(supports, loads):
        st = ProjectState(
            length=L, materials=[matB], selected_material_id="mb",
            cross_section=rect, section_segments=[SectionSegment(0.0, L, rect, None)],
            supports=supports, load_cases=[LoadCase("lc", "LC", False)], loads=loads,
            load_combinations=[LoadCombination(id="c", name="c", factors={l.id: 1.0 for l in loads})],
            additional_factor=1.0)
        return buckling_eigen_check(st, solve_beam(st))

    # kloub-kloub: osový tlak koncovou silou (rolna nechá u volné) → μ=1
    bk = run([Support("a", 0, "pin", 0), Support("b", L, "roller", 0)],
             [Load("f", "point_force", "F", "lc", x=L, Fx=-P)])
    assert _rel(bk.mu_eff, 1.0) < 1e-2 and _rel(bk.P_cr, P_euler) < 1e-2

    # vetknutý-volný (konzola): μ=2, P_cr = Euler/4
    bk = run([Support("a", 0, "fixed", 0)],
             [Load("f", "point_force", "F", "lc", x=L, Fx=-P)])
    assert _rel(bk.mu_eff, 2.0) < 1e-2 and _rel(bk.P_cr, 0.25 * P_euler) < 1e-2

    # vetknutý-vetknutý: tlak z omezené teplotní dilatace (oba konce drží u) → μ=0.5
    th = Load("t", "thermal", "T", "lc"); th.x1 = 0; th.x2 = L; th.dT = 50.0
    bk = run([Support("a", 0, "fixed", 0), Support("b", L, "fixed", 0)], [th])
    assert _rel(bk.mu_eff, 0.5) < 1e-2 and _rel(bk.P_cr, 4.0 * P_euler) < 1e-2

    # vetknutý-kloub (propped): μ=0.699 – TRANSCENDENTNÍ řešení (tan(kL)=kL),
    # smíšená OP, ověří eigen-solver na netriviálním případu. P_cr = π²EI/(0.699L)²
    bk = run([Support("a", 0, "fixed", 0), Support("b", L, "roller", 0)],
             [Load("f", "point_force", "F", "lc", x=L, Fx=-P)])
    assert _rel(bk.mu_eff, 0.6992) < 5e-3 and _rel(bk.P_cr, P_euler / 0.6992**2) < 5e-3

    # čistý tah → žádné vybočení (None)
    assert run([Support("a", 0, "pin", 0), Support("b", L, "roller", 0)],
               [Load("f", "point_force", "F", "lc", x=L, Fx=+P)]) is None


def test_thermal_gradient_bending():
    """Teplotní gradient (fáze 2): rozdíl teplot horní−dolní vlákno → křivost
    κ=α·ΔT_grad/h. Prostě podepřený se volně prohne (M≈0, δ=κ·L²/8, bez pnutí);
    vetknutý-vetknutý má konstantní M=E·I·κ (napětí od zabráněné křivosti)."""
    E, alpha, b, h, L, dTg = 210000.0, 12e-6, 40.0, 80.0, 1000.0, 50.0
    matT = Material("mt", "St", E=E, G=81000.0, nu=0.3, rho=7.85,
                    Re=235.0, Rm=360.0, alpha=alpha)
    rect = CrossSectionDef(type="rectangle", params={"b": b, "h": h})
    I = b * h**3 / 12.0
    kth = alpha * dTg / h
    M_exp = E * I * kth
    d_exp = kth * L**2 / 8.0

    def mk(supports):
        th = Load("t", "thermal", "dT", "lc"); th.x1 = 0; th.x2 = L
        th.dT = 0.0; th.dT_grad = dTg
        return ProjectState(
            length=L, materials=[matT], selected_material_id="mt",
            cross_section=rect, section_segments=[SectionSegment(0.0, L, rect, None)],
            supports=supports, load_cases=[LoadCase("lc", "LC", False)], loads=[th],
            load_combinations=[LoadCombination(id="c", name="c", factors={"t": 1.0})],
            additional_factor=1.0)

    # prostě podepřený: volná křivost → M≈0, průhyb δ=κ·L²/8
    r_ss = solve_beam(mk([Support("a", 0, "pin", 0), Support("b", L, "roller", 0)]))
    assert max(abs(p.M) for p in r_ss.points) < 1e-3 * M_exp       # bez pnutí
    w_mid = [p.w for p in r_ss.points if abs(p.x - L/2) < 6][0]
    assert _rel(abs(w_mid), d_exp) < 5e-3                          # tepelná klenba
    # vetknutý-vetknutý: zabráněná křivost → konstantní M=E·I·κ, průhyb ~0
    r_ff = solve_beam(mk([Support("a", 0, "fixed", 0), Support("b", L, "fixed", 0)]))
    Ms = [p.M for p in r_ff.points]
    assert _rel(abs(Ms[len(Ms)//2]), M_exp) < 1e-6                 # M = E·I·κ
    assert max(Ms) - min(Ms) < 1e-3 * M_exp                       # konstantní po délce
    assert max(abs(p.w) for p in r_ff.points) < 1e-6              # bez průhybu


if __name__ == "__main__":
    tests = [v for k, v in sorted(globals().items()) if k.startswith("test_")]
    passed = failed = 0
    for t in tests:
        try:
            t()
            print(f"  PASS  {t.__name__}")
            passed += 1
        except AssertionError as e:
            print(f"  FAIL  {t.__name__}: {e}")
            failed += 1
        except Exception as e:
            print(f"  ERROR {t.__name__}: {e!r}")
            failed += 1
    print(f"\n{passed} passed, {failed} failed")
    sys.exit(1 if failed else 0)


# ═══════════════════════════════════════════════════════════════
#  B2 – tagging elementů materiálem (podpora variabilního G FEM)
# ═══════════════════════════════════════════════════════════════

def test_b2_element_region_tagging():
    """Element se přiřadí té oblasti, do které padne jeho těžiště."""
    from beamer.composite_fem import assign_element_regions
    nodes = [(0.2, 0.5), (0.4, 0.5), (0.3, 0.8),      # trojúhelník v oblasti A
             (1.2, 0.5), (1.4, 0.5), (1.3, 0.8)]      # trojúhelník v oblasti B
    elements = [[0, 1, 2], [3, 4, 5]]
    regions = [(([(0, 0), (1, 0), (1, 1), (0, 1)], []), "A"),
               (([(1, 0), (2, 0), (2, 1), (1, 1)], []), "B")]
    idx, payloads = assign_element_regions(nodes, elements, regions)
    assert idx == [0, 1]
    assert payloads[idx[0]] == "A" and payloads[idx[1]] == "B"


def test_b2_tagging_respects_holes():
    """Těžiště v díře oblasti → element se do ní nezařadí."""
    from beamer.composite_fem import assign_element_regions
    nodes = [(4.9, 5.0), (5.1, 5.0), (5.0, 5.2)]      # střed velkého čtverce s dírou
    elements = [[0, 1, 2]]
    outer = [(0, 0), (10, 0), (10, 10), (0, 10)]
    hole = [(4, 4), (6, 4), (6, 6), (4, 6)]           # díra kolem středu
    regions = [((outer, [hole]), "S")]
    idx, _ = assign_element_regions(nodes, elements, regions)
    assert idx == [-1]                                 # v díře → nezařazeno


def test_b2_variable_G_reduces_and_scales():
    """Variabilní-G warping: konstantní G=1 dá IDENTICKÉ J jako dosud (regrese),
    konstantní G=2 dá GJ_eff = 2·J (správné škálování)."""
    from beamer import _fem
    _fem.set_element_order("T6")
    outer = [(-50, -30), (50, -30), (50, 30), (-50, 30)]   # obdélník 100×60
    nodes, elements = _fem.triangulate_section(outer, None)
    g = _fem.compute_geometric_properties(nodes, elements)
    cy, cz, Ixx, Iyy = g['cy'], g['cz'], g['Ixx_c'], g['Iyy_c']
    om0, K0 = _fem.solve_warping_function(nodes, elements, cy, cz)
    J0 = _fem.compute_torsion_constant(nodes, elements, om0, K0, Ixx, Iyy)
    G1 = [1.0]*len(elements)
    om1, K1 = _fem.solve_warping_function(nodes, elements, cy, cz, elem_G=G1)
    J1 = _fem.compute_torsion_constant(nodes, elements, om1, K1, Ixx, Iyy, elem_G=G1, cy=cy, cz=cz)
    assert _rel(J1, J0) < 1e-9
    G2 = [2.0]*len(elements)
    om2, K2 = _fem.solve_warping_function(nodes, elements, cy, cz, elem_G=G2)
    J2 = _fem.compute_torsion_constant(nodes, elements, om2, K2, Ixx, Iyy, elem_G=G2, cy=cy, cz=cz)
    assert _rel(J2, 2.0*J0) < 1e-9


def _composite_state(parts, sections, materials):
    from beamer.model import Property
    prop = Property(id="p", pid=1, name="c", composite_parts=parts)
    st = ProjectState(length=100, sections=sections, materials=materials,
                      selected_material_id=materials[0].id, properties=[prop])
    return st, prop


def test_b2_composite_GJ_single_material_reduces():
    """(GJ)_eff složeného ze STEJNÉHO materiálu = G·J geometrického průřezu."""
    from beamer.composite_fem import composite_torsion_GJ
    from beamer import _fem
    mat = Material("a", "Ocel", E=210000, G=81000, nu=0.3, rho=7.85, Re=235, Rm=360)
    r = CrossSectionDef(type="rectangle", params={"b": 40, "h": 20}, id="r", name="R")
    st, prop = _composite_state(
        [{"section_id": "r", "material_id": "a", "dy": 0, "dz": 10, "angle": 0},
         {"section_id": "r", "material_id": "a", "dy": 0, "dz": -10, "angle": 0}],
        [r], [mat])
    GJ = composite_torsion_GJ(st, prop)
    # geometrická torze celku 40×40 (dvě 40×20 na sobě) přes stejný mesh
    _fem.set_element_order("T6")
    outer = [(-20, -20), (20, -20), (20, 20), (-20, 20)]
    nodes, elements = _fem.triangulate_section(outer, None)
    g = _fem.compute_geometric_properties(nodes, elements)
    om, K = _fem.solve_warping_function(nodes, elements, g["cy"], g["cz"])
    J = _fem.compute_torsion_constant(nodes, elements, om, K, g["Ixx_c"], g["Iyy_c"])
    assert _rel(GJ, 81000.0 * J) < 5e-3


def test_b2_composite_GJ_concentric_analytic():
    """Souosá tyč v trubce (různé G, ω=0): (GJ)_eff = G_tyč·J_tyč + G_trubka·J_trubka."""
    from beamer.composite_fem import composite_torsion_GJ
    steel = Material("st", "Ocel", E=210000, G=81000, nu=0.3, rho=7.85, Re=235, Rm=360)
    alu = Material("al", "Hliník", E=70000, G=27000, nu=0.33, rho=2.7, Re=200, Rm=300)
    rod = CrossSectionDef(type="circle", params={"D": 20}, id="rod", name="tyč")
    tube = CrossSectionDef(type="tube", params={"Do": 40, "t": 10}, id="tube", name="trubka")
    st, prop = _composite_state(
        [{"section_id": "rod", "material_id": "st", "dy": 0, "dz": 0, "angle": 0},
         {"section_id": "tube", "material_id": "al", "dy": 0, "dz": 0, "angle": 0}],
        [rod, tube], [steel, alu])
    GJ = composite_torsion_GJ(st, prop)
    Jrod = math.pi * 10**4 / 2
    Jtube = math.pi * (20**4 - 10**4) / 2
    GJ_anal = 81000 * Jrod + 27000 * Jtube
    assert _rel(GJ, GJ_anal) < 1e-2      # polygonizace kruhu ~0,3 %


def test_xval_composite_GJ_vs_sectionproperties():
    """Vázaný variabilní-G warping (nesymetrický spoj, různé G) vs sectionproperties.
    E_sp := G_mat → sectionproperties get_ej váží stejným modulem jako naše GJ."""
    import pytest
    pytest.importorskip("sectionproperties")
    from sectionproperties.pre.library import rectangular_section as rs
    from sectionproperties.pre import Material as SPMat
    from sectionproperties.analysis import Section
    from beamer.composite_fem import composite_torsion_GJ
    G1, G2 = 81000.0, 27000.0
    m1 = Material("a", "Tuhý", E=2*G1*1.3, G=G1, nu=0.3, rho=7.85, Re=235, Rm=360)
    m2 = Material("b", "Měkký", E=2*G2*1.3, G=G2, nu=0.3, rho=2.7, Re=200, Rm=300)
    r = CrossSectionDef(type="rectangle", params={"b": 40, "h": 20}, id="r", name="R")
    st, prop = _composite_state(
        [{"section_id": "r", "material_id": "a", "dy": 0, "dz": 10, "angle": 0},
         {"section_id": "r", "material_id": "b", "dy": 0, "dz": -10, "angle": 0}],
        [r], [m1, m2])
    GJ = composite_torsion_GJ(st, prop)
    a_sp = SPMat("a", G1, 0.3, 235, 7.85e-6, "grey")
    b_sp = SPMat("b", G2, 0.3, 200, 2.7e-6, "blue")
    geom = (rs(d=20, b=40, material=a_sp).shift_section(x_offset=-20, y_offset=0)
            + rs(d=20, b=40, material=b_sp).shift_section(x_offset=-20, y_offset=-20))
    geom.create_mesh(mesh_sizes=[1.0])
    s = Section(geom); s.calculate_geometric_properties(); s.calculate_warping_properties()
    assert _rel(GJ, s.get_ej()) < 2e-2


def test_b2_composite_von_mises_torsion():
    """Plný von Mises složeného: čistý ohyb → τ=0, mises=σ; čistá torze → σ=0,
    mises=√3·τ a τ per materiál ≈ Gᵢ·θ'·r (souosá tyč v trubce)."""
    from beamer.composite_fem import composite_stress_field, composite_torsion_GJ
    steel = Material("st", "Ocel", E=210000, G=81000, nu=0.3, rho=7.85, Re=235, Rm=360)
    alu = Material("al", "Hliník", E=70000, G=27000, nu=0.33, rho=2.7, Re=200, Rm=300)
    rod = CrossSectionDef(type="circle", params={"D": 20}, id="rod", name="tyč")
    tube = CrossSectionDef(type="tube", params={"Do": 40, "t": 10}, id="tube", name="trubka")
    st, prop = _composite_state(
        [{"section_id": "rod", "material_id": "st", "dy": 0, "dz": 0, "angle": 0},
         {"section_id": "tube", "material_id": "al", "dy": 0, "dz": 0, "angle": 0}],
        [rod, tube], [steel, alu])
    # čistý ohyb → τ=0, mises=σ
    fb = composite_stress_field(st, prop, 0.0, 1.0e6, 0.0)
    for a in fb["materials"].values():
        assert a["tau_max"] < 1e-6
        assert _rel(a["mises_max"], a["sigma_max"]) < 1e-9
    # čistá torze → σ=0, mises=√3·τ, τ per materiál ≈ Gᵢ·θ'·r
    Mk = 1.0e6
    GJ = composite_torsion_GJ(st, prop)
    thp = Mk / GJ
    ft = composite_stress_field(st, prop, 0.0, 0.0, Mk)
    assert ft["materials"]["st"]["sigma_max"] < 1e-6
    assert _rel(ft["materials"]["al"]["mises_max"],
                math.sqrt(3) * ft["materials"]["al"]["tau_max"]) < 1e-9
    assert _rel(ft["materials"]["al"]["tau_max"], 27000 * thp * 20) < 3e-2   # trubka r=20
    assert _rel(ft["materials"]["st"]["tau_max"], 81000 * thp * 10) < 6e-2   # tyč r=10 (rozhraní)


def test_b2_composite_transverse_shear_reduces():
    """Transverzální smyk τ_V složeného ze STEJNÉHO materiálu = Žuravskij
    obdélníku: τ_max = 1.5·V/A (E-vážený Žuravskij se redukuje na klasický)."""
    from beamer.composite_fem import composite_stress_field
    steel = Material("st", "Ocel", E=210000, G=81000, nu=0.3, rho=7.85, Re=235, Rm=360)
    r = CrossSectionDef(type="rectangle", params={"b": 40, "h": 20}, id="r", name="R")
    st, prop = _composite_state(
        [{"section_id": "r", "material_id": "st", "dy": 0, "dz": 10, "angle": 0},
         {"section_id": "r", "material_id": "st", "dy": 0, "dz": -10, "angle": 0}],
        [r], [steel])
    V = 1000.0
    f = composite_stress_field(st, prop, 0.0, 0.0, 0.0, V)
    tauV = f["materials"]["st"]["tau_max"]
    assert _rel(tauV, 1.5 * V / (40 * 40)) < 1e-3      # 40×40, τ_max = 1.5 V/A


def _bimetal_state():
    """Nesymetrický bimetal: ocel nahoře (z 0..20), hliník dole (z −20..0)."""
    steel = Material("st", "Ocel", E=210000, G=81000, nu=0.3, rho=7.85, Re=235, Rm=360)
    alu = Material("al", "Hliník", E=70000, G=27000, nu=0.33, rho=2.7, Re=200, Rm=300)
    r = CrossSectionDef(type="rectangle", params={"b": 40, "h": 20}, id="r", name="R")
    return _composite_state(
        [{"section_id": "r", "material_id": "st", "dy": 0, "dz": 10, "angle": 0},
         {"section_id": "r", "material_id": "al", "dy": 0, "dz": -10, "angle": 0}],
        [r], [steel, alu])


def test_composite_bending_sign_with_axial():
    """Znaménko ohybu u kompozitu (N1): M+ = sagging = tah DOLE (konvence
    solveru). Pro N+M na NEsymetrickém bimetalu musí σ per materiál odpovídat
    σ = E·(N/EA − M·(z−z_NA)/EIy); s obráceným znaménkem by hliník (dole)
    vyšel o ~36 % nekonzervativně."""
    from beamer.composite import composite_stress, composite_weighted
    from beamer.composite_fem import composite_stress_field
    st, prop = _bimetal_state()
    w = composite_weighted(st, prop)
    N, M = 50e3, 1.0e6

    def sig_ok(E, z):
        return E * (N / w.EA - M * (z - w.z_NA) / w.EIy)

    expect = {"Ocel": max(abs(sig_ok(210000, z)) for z in (0.0, 20.0)),
              "Hliník": max(abs(sig_ok(70000, z)) for z in (-20.0, 0.0))}
    by = {r_["material"]: r_ for r_ in composite_stress(st, prop, N, M)}
    for nm in expect:
        assert _rel(by[nm]["sigma_max"], expect[nm]) < 1e-9
    f = composite_stress_field(st, prop, N, M, 0.0, 0.0)
    for a in f["materials"].values():
        assert _rel(a["sigma_max"], expect[a["material"]]) < 1e-9
    # hliník (dole, tah se přičítá k tahu) musí nést víc než samotné N/EA
    assert by["Hliník"]["sigma_max"] > 70000 * (N / w.EA)


def test_composite_field_sigma_reaches_fiber():
    """σ z FEM pole (N3): vyhodnocení v rozích elementů dosáhne přesné krajní
    vlákno – žádné podhodnocení proti B1 (těžiště elementu by dalo ~2 % míň)."""
    from beamer.composite import composite_stress
    from beamer.composite_fem import composite_stress_field
    st, prop = _bimetal_state()
    by = {r_["material"]: r_ for r_ in composite_stress(st, prop, 0.0, 1.0e6)}
    f = composite_stress_field(st, prop, 0.0, 1.0e6, 0.0, 0.0)
    for a in f["materials"].values():
        assert _rel(a["sigma_max"], by[a["material"]]["sigma_max"]) < 1e-9


def test_composite_mesh_cache_sees_polygon_edit():
    """Signatura cache kompozitní sítě (N4): editace polygon_points knihovního
    profilu musí zneplatnit cache – GJ se přepočítá (dřív zůstal starý)."""
    from beamer.composite_fem import composite_torsion_GJ
    steel = Material("st", "Ocel", E=210000, G=81000, nu=0.3, rho=7.85, Re=235, Rm=360)
    alu = Material("al", "Hliník", E=70000, G=27000, nu=0.33, rho=2.7, Re=200, Rm=300)
    poly = CrossSectionDef(type="polygon", id="pg", name="P", polygon_points=[
        {"y": -20, "z": -10}, {"y": 20, "z": -10},
        {"y": 20, "z": 10}, {"y": -20, "z": 10}])
    r = CrossSectionDef(type="rectangle", params={"b": 40, "h": 20}, id="r", name="R")
    st, prop = _composite_state(
        [{"section_id": "pg", "material_id": "st", "dy": 0, "dz": 10, "angle": 0},
         {"section_id": "r", "material_id": "al", "dy": 0, "dz": -10, "angle": 0}],
        [poly, r], [steel, alu])
    GJ1 = composite_torsion_GJ(st, prop)
    # editace geometrie polygonu (2× šířka) → GJ se MUSÍ změnit
    poly.polygon_points = [{"y": -40, "z": -10}, {"y": 40, "z": -10},
                           {"y": 40, "z": 10}, {"y": -40, "z": 10}]
    GJ2 = composite_torsion_GJ(st, prop)
    assert GJ1 is not None and GJ2 is not None
    assert abs(GJ2 - GJ1) / GJ1 > 0.1


def test_composite_combined_sigma_red_mode():
    """Režim σ_red „combined" (N6) působí i na kompozit: σ_red per materiál
    = √(σ_max²+3·τ_max²) ze špiček; bez fallbacku je b1_fallback False."""
    import math
    from beamer.composite import composite_assess
    st, prop = _bimetal_state()
    st.sigma_red_mode = "combined"
    ca = composite_assess(st, prop, N=10e3, M=5e5, basis="min", Mk=2e5, V=1e3)
    assert ca is not None and not ca.get("b1_fallback")
    assert ca.get("sigma_red_combined") is True
    for m in ca["materials"]:
        assert _rel(m["mises_max"],
                    math.sqrt(m["sigma_max"]**2 + 3*m["tau_max"]**2)) < 1e-9


def test_material_library_order_and_tolerant_load(tmp_path, monkeypatch):
    """Knihovna materiálů: save_materials zachovává pořadí (organizace „oceli
    k sobě") a čtení toleruje neznámé klíče (starší/novější formát nesmí
    položku zahodit)."""
    from beamer import library
    monkeypatch.setattr(library, "_LOCAL_DIR", str(tmp_path))
    a = Material("a", "Ocel S355", 210000, 81000, 0.3, 7.85, 355, 490)
    b = Material("b", "Al 6061", 69000, 26000, 0.33, 2.7, 240, 290)
    assert library.save_materials([a, b])
    assert [m.name for m in library.load_materials()] == ["Ocel S355", "Al 6061"]
    library.save_materials([b, a])            # přeuspořádání
    assert [m.name for m in library.load_materials()] == ["Al 6061", "Ocel S355"]
    import json, os
    p = os.path.join(str(tmp_path), "materials.json")
    with open(p, encoding="utf-8") as f:
        data = json.load(f)
    data[0]["budouci_pole"] = 123             # neznámý klíč
    with open(p, "w", encoding="utf-8") as f:
        json.dump(data, f)
    assert [m.name for m in library.load_materials()] == ["Al 6061", "Ocel S355"]


def test_live_material_combo_import_and_dedup(tmp_path, monkeypatch):
    """Živá knihovna v combech: knihovní položka se nabízí, výběr ji ZKOPÍRUJE
    do projektu (vč. α/Fcy), druhý výběr vrátí existující kopii (dedup) a
    zkopírovaná položka se v nabídce už neduplikuje."""
    pytest.importorskip("PySide6")
    import os
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    from PySide6.QtWidgets import QApplication
    _ = QApplication.instance() or QApplication([])
    from beamer import library
    monkeypatch.setattr(library, "_LOCAL_DIR", str(tmp_path))
    library.save_materials([
        Material("l1", "Ocel S355", 210000, 81000, 0.3, 7.85, 355, 490,
                 alpha=12e-6, Fcy=340.0),
    ])
    from beamer.gui.widgets import fill_material_combo, resolve_material_choice
    from beamer.gui.spin import NoWheelComboBox
    st = ProjectState(materials=[Material("m1", "Projektový", 70000, 27000,
                                          0.33, 2.7, 200, 300)],
                      selected_material_id="m1")
    cb = NoWheelComboBox()
    fill_material_combo(cb, st, "m1")
    datas = [cb.itemData(i) for i in range(cb.count())]
    libs = [d for d in datas if isinstance(d, tuple)]
    assert len(libs) == 1 and libs[0][1].name == "Ocel S355"
    assert cb.currentData() == "m1"           # aktuální výběr zachován

    mid, created = resolve_material_choice(st, libs[0])
    assert created and len(st.materials) == 2
    m = next(mm for mm in st.materials if mm.id == mid)
    assert m.name == "Ocel S355" and m.Fcy == 340.0 and m.alpha == 12e-6
    mid2, created2 = resolve_material_choice(st, libs[0])
    assert mid2 == mid and not created2       # dedup: žádná další kopie

    cb2 = NoWheelComboBox()                   # po importu už není v nabídce 2×
    fill_material_combo(cb2, st, mid)
    libs2 = [cb2.itemData(i) for i in range(cb2.count())
             if isinstance(cb2.itemData(i), tuple)]
    assert not libs2
    assert cb2.currentData() == mid


def test_profile_library_order_roundtrip(tmp_path, monkeypatch):
    """Knihovna profilů: save_profiles zachovává pořadí (organizace „trubky
    k sobě") a save_profile (upsert) pořadí nerozbije."""
    from beamer import library
    monkeypatch.setattr(library, "_LOCAL_DIR", str(tmp_path))
    a = CrossSectionDef(type="rectangle", params={"b": 40, "h": 20})
    b = CrossSectionDef(type="tube", params={"Do": 60, "t": 4})
    assert library.save_profiles([("Obdelnik", a), ("Trubka", b)])
    assert [n for n, _ in library.load_profiles()] == ["Obdelnik", "Trubka"]
    library.save_profiles([("Trubka", b), ("Obdelnik", a)])
    assert [n for n, _ in library.load_profiles()] == ["Trubka", "Obdelnik"]
    library.save_profile("Obdelnik", a)       # upsert existujícího
    assert [n for n, _ in library.load_profiles()] == ["Trubka", "Obdelnik"]


def test_live_section_combo_import_and_dedup(tmp_path, monkeypatch):
    """Živá knihovna profilů ve výběru průřezu: knihovní profil se nabízí,
    výběr ho zkopíruje do projektových průřezů, dedup nevytváří další kopie,
    inline volba (None) zůstává funkční."""
    pytest.importorskip("PySide6")
    import os
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    from PySide6.QtWidgets import QApplication
    _ = QApplication.instance() or QApplication([])
    from beamer import library
    monkeypatch.setattr(library, "_LOCAL_DIR", str(tmp_path))
    library.save_profiles([("Trubka 60x4",
                            CrossSectionDef(type="tube", params={"Do": 60, "t": 4}))])
    from beamer.gui.widgets import (fill_section_combo, resolve_section_choice,
                                    _HDR)
    from beamer.gui.spin import NoWheelComboBox
    st = ProjectState(sections=[CrossSectionDef(type="rectangle",
                                                params={"b": 100, "h": 200},
                                                id="s1", name="R1")])
    cb = NoWheelComboBox()
    fill_section_combo(cb, st, None)
    datas = [cb.itemData(i) for i in range(cb.count())]
    assert cb.currentData() is None            # inline výchozí
    libs = [d for d in datas if isinstance(d, tuple)]
    assert len(libs) == 1 and libs[0][1] == "Trubka 60x4"
    assert _HDR in datas                       # nadpis skupiny přítomen

    sid, created = resolve_section_choice(st, libs[0])
    assert created and len(st.sections) == 2
    sec = next(s for s in st.sections if s.id == sid)
    assert sec.name == "Trubka 60x4" and sec.type == "tube"
    sid2, created2 = resolve_section_choice(st, libs[0])
    assert sid2 == sid and not created2        # dedup

    cb2 = NoWheelComboBox()                    # po importu už není v nabídce 2×
    fill_section_combo(cb2, st, sid)
    assert not [cb2.itemData(i) for i in range(cb2.count())
                if isinstance(cb2.itemData(i), tuple)]
    assert cb2.currentData() == sid
    # inline volba dál funguje
    assert resolve_section_choice(st, None) == (None, False)


def test_profile_library_roundtrip_preserves_all_fields(tmp_path, monkeypatch):
    """P1 (review 1.34): knihovna profilů nesmí při načtení ztratit shapes,
    rotation ani Body.material_id – konstrukční/natočený/kompozitní profil
    projde uložením a načtením beze změny definice."""
    from beamer import library
    from beamer.model import Body
    monkeypatch.setattr(library, "_LOCAL_DIR", str(tmp_path))
    sdef = CrossSectionDef(
        type="construction",
        shapes=[{"kind": "rect", "op": "add", "y": 0, "z": 0, "b": 40, "h": 20}],
        rotation=30.0,
        bodies=[Body(points=[{"y": 0, "z": 0}, {"y": 10, "z": 0},
                             {"y": 10, "z": 10}],
                     holes=[], material_id="mat_x")])
    library.save_profiles([("Konstrukcni", sdef)])
    (name, back), = library.load_profiles()
    assert name == "Konstrukcni"
    assert back.type == "construction"
    assert back.rotation == 30.0
    assert back.shapes == sdef.shapes
    assert back.bodies and back.bodies[0].material_id == "mat_x"
    assert back.bodies[0].points == sdef.bodies[0].points


def test_composite_add_starts_from_library_in_empty_project(tmp_path, monkeypatch):
    """P2 (review 1.34): nový projekt bez projektových průřezů – „Přidat
    profil" ve skladbě si vezme první použitelný profil z KNIHOVNY (zkopíruje
    do projektu); „direct" v knihovně se přeskočí."""
    pytest.importorskip("PySide6")
    import os
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    from PySide6.QtWidgets import QApplication
    _ = QApplication.instance() or QApplication([])
    from beamer import library
    monkeypatch.setattr(library, "_LOCAL_DIR", str(tmp_path))
    library.save_profiles([
        ("Direct X", CrossSectionDef(type="direct",
                                     params={"A": 1e3, "Iy": 1e6, "Iz": 1e6,
                                             "IT": 1e6})),
        ("Trubka 60x4", CrossSectionDef(type="tube", params={"Do": 60, "t": 4})),
    ])
    from beamer.model import Property
    from beamer.gui.composite_dialog import CompositeEditorDialog
    st = ProjectState(materials=[MAT], selected_material_id=MAT.id)
    assert not st.sections                     # prázdný projekt
    p = Property(id="p", pid=1, name="komp", composite_parts=[])
    dlg = CompositeEditorDialog(st, p)
    dlg._add()
    assert len(p.composite_parts) == 1         # část přidána z knihovny
    sec = next(s for s in st.sections
               if s.id == p.composite_parts[0]["section_id"])
    assert sec.type == "tube" and sec.name == "Trubka 60x4"   # direct přeskočen


def test_new_load_is_registered_in_combinations():
    """Nově přidané zatížení nesmí být TICHÁ NULA.

    Jakmile jsou faktory kombinace klíčované podle id zatížení (což nastane po
    otevření projektu ze souboru nebo po otevření Load Case Builderu), nové
    zatížení v mapě chybí a solver mu dá faktor 0 – uživatel zadá sílu a „nic
    se nepočítá". `register_load_in_combinations` to napraví; vědomé vyřazení
    (explicitní 0.0) se přitom nesmí přepsat."""
    from beamer.defaults import create_default_state
    from beamer.model import (register_load_in_combinations, combination_factor,
                              migrate_combinations_to_loads, new_id)
    from beamer.solver import _load_multiplier

    st = create_default_state()
    migrate_combinations_to_loads(st)          # = otevření projektu ze souboru
    comb = st.active_combination()
    assert set(comb.factors) == {st.loads[0].id}      # klíčováno dle zatížení

    new = Load(new_id("load"), "point_force", "Nová síla", st.load_cases[0].id)
    new.x, new.Fz = 1000.0, -5000.0
    st.loads.append(new)
    assert _load_multiplier(st, new) == 0.0           # bez registrace tichá nula
    register_load_in_combinations(st, new)
    assert _load_multiplier(st, new) == 1.0

    # vědomé vyřazení se nepřepisuje
    ex = st.loads[0]
    comb.factors[ex.id] = 0.0
    register_load_in_combinations(st, ex)
    assert comb.factors[ex.id] == 0.0
    assert combination_factor(comb, ex) == 0.0

    # zatížení se skutečně projeví ve výsledku
    m_before = max(abs(p.M) for p in solve_beam(st).points)
    comb.factors[ex.id] = 1.0
    assert max(abs(p.M) for p in solve_beam(st).points) > m_before


def test_axial_displacement_matches_analytic():
    """Osový posun u(x): tažený vetknutý prut ΔL = N·L/(E·A), lineární průběh;
    volná teplotní dilatace ΔL = α·ΔT·L při nulové osové síle."""
    L, P = 2000.0, 50000.0
    rect = CrossSectionDef(type="rectangle", params={"b": 40, "h": 20})
    st = ProjectState(
        length=L, supports=[Support("a", 0, "fixed", 0)],
        load_cases=[LoadCase("lc", "LC", False)],
        load_combinations=[LoadCombination("c", "C", {"lc": 1.0})],
        loads=[Load("f", "point_force", "F", "lc", x=L, Fx=P)],
        materials=[MAT], selected_material_id=MAT.id, cross_section=rect,
        section_segments=[SectionSegment(0.0, L, rect)],
        additional_factor=1.0, selected_active_combination_id="c")
    res = solve_beam(st)
    A = build_section(rect, fem=False).A
    dL = P * L / (MAT.E * A)
    assert _rel(res.points[-1].u, dL) < 1e-9
    assert abs(res.points[0].u) < 1e-9
    mid = res.points[len(res.points) // 2]
    assert _rel(mid.u, dL * mid.x / L) < 1e-6          # lineární průběh

    # volná teplotní dilatace: N = 0, přesto se prut prodlouží
    th = Load("t", "thermal", "dT", "lc"); th.x1 = 0; th.x2 = L; th.dT = 100.0
    st.loads = [th]
    st.load_combinations[0].factors = {"t": 1.0}
    res2 = solve_beam(st)
    assert abs(res2.points[0].N) < 1e-6
    assert _rel(res2.points[-1].u, MAT.alpha * 100.0 * L) < 1e-9


def test_plastic_factor_shear_interaction():
    """α_pl s interakcí smyku: plná hodnota při čistém ohybu, parabolický pokles
    k 1.0 na hranici R_v = τ_V/τ_dov = 0.25, nad ní se neuplatní. Osová síla,
    torze a biaxiál α vypínají úplně."""
    from beamer.analysis import (_plastic_capacity_factor, shear_allowable,
                                 PLASTIC_SHEAR_LIMIT)
    a, M = 1.15, 1.0e6
    ta = shear_allowable(483.0, None)
    assert _rel(ta, 483.0 / math.sqrt(3.0)) < 1e-12
    assert shear_allowable(483.0, 290.0) == 290.0      # Fsu má přednost

    # čistý ohyb → plná α
    assert _plastic_capacity_factor(a, True, 0, 0, M, 0) == a
    # monotónní pokles k 1.0 na hranici
    prev = a
    for frac in (0.05, 0.1, 0.2, 0.249):
        cur = _plastic_capacity_factor(a, True, 0, 1.0, M, 0,
                                       tau_v=frac * ta, tau_allow=ta)
        assert 1.0 < cur <= prev
        prev = cur
    assert _plastic_capacity_factor(a, True, 0, 1.0, M, 0,
                                    tau_v=PLASTIC_SHEAR_LIMIT * ta,
                                    tau_allow=ta) == 1.0
    assert _plastic_capacity_factor(a, True, 0, 1.0, M, 0,
                                    tau_v=0.9 * ta, tau_allow=ta) == 1.0
    # osová síla / torze / biaxiál vypínají
    for kw in ({"N": 1e4}, {"Mk": 1e5}, {"Mz": 1e5}, {"Vy": 1e3}):
        kwargs = {"N": 0, "V": 1.0, "M": M, "Mk": 0}
        kwargs.update(kw)
        assert _plastic_capacity_factor(
            a, True, kwargs["N"], kwargs["V"], kwargs["M"], kwargs["Mk"],
            Mz=kwargs.get("Mz", 0.0), Vy=kwargs.get("Vy", 0.0),
            tau_v=0.0, tau_allow=ta) == 1.0
    # vypnutá plasticita / neznámá smyková únosnost
    assert _plastic_capacity_factor(a, False, 0, 0, M, 0) == 1.0
    assert _plastic_capacity_factor(a, True, 0, 1.0, M, 0,
                                    tau_v=1.0, tau_allow=0.0) == 1.0


def test_plasticity_raises_ultimate_rf_along_beam():
    """Zapnutí plasticity se MUSÍ projevit na RF_ultimate podél nosníku (dříve
    ji gate vypínal při jakémkoli smyku, takže funkce byla fakticky mrtvá)."""
    from beamer.defaults import create_default_state
    from beamer.analysis import reserves_along_beam
    st = create_default_state()
    st.rf_basis = "ultimate"
    res = solve_beam(st)
    st.plasticity_enabled = False
    rf_off = min(mm.RF for mm in reserves_along_beam(res, st, n_stations=40))
    st.plasticity_enabled = True
    rf_on = min(mm.RF for mm in reserves_along_beam(res, st, n_stations=40))
    assert rf_on > rf_off * 1.05
    # při bázi min řídí mez kluzu → α_pl se (správně) neprojeví
    st.rf_basis = "min"
    st.plasticity_enabled = False
    a = min(mm.RF for mm in reserves_along_beam(res, st, n_stations=40))
    st.plasticity_enabled = True
    b = min(mm.RF for mm in reserves_along_beam(res, st, n_stations=40))
    assert _rel(a, b) < 1e-12
